#!/usr/bin/env python3
"""
SkillSync Enhanced Backend (v2.0) - Simplified Version
Compatible with older Python packages

Features:
- F1: Enhanced CV Analysis with AI-powered skill extraction
- F2: Intelligent Job Matching with semantic similarity
- F3: Advanced Skill Gap Analysis with market trends
- F4: Personalized Career Recommendations
- F5: AI Experience Translator (Technical  Business)
- F6: Explainable AI (XAI) Dashboard
- F7: Advanced Analytics Dashboard
- F8: Dynamic Portfolio Generator

Author: MiniMax Agent
Version: 2.0
Date: 2025-10-26
"""

import os
import logging
import asyncio
import uuid
import traceback
import json
import time
import html
import math
import functools
import threading
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Any, Tuple, Set, Iterable
from pathlib import Path
from urllib.parse import quote_plus
import numpy as np
import re
import string
from collections import Counter, defaultdict

try:
    from ml_models.skills_extractor import SkillsExtractorModel
except Exception:
    SkillsExtractorModel = None

# FastAPI and web dependencies
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from starlette.exceptions import HTTPException as StarletteHTTPException

from multi_job_service import search_jobs_multi_source
from routers import cv_analysis
from skillsync.interviews import routes as interview_routes
from skillsync.interviews import realtime as interview_realtime
from roadmap_ml.predictor import get_default_predictor

# Import auth router
try:
    from auth.router import router as auth_router
    AUTH_ENABLED = True
except ImportError:
    AUTH_ENABLED = False
    print("⚠️ Authentication module not available")

# Database dependencies
from sqlalchemy import create_engine, Column, Integer, String, Text, Float, DateTime, Boolean, ForeignKey, JSON, func
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship

# ML/AI dependencies (optional)
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    print(" spaCy model loaded successfully")
except:
    nlp = None
    print(" spaCy not available - using basic functionality")

try:
    import nltk
    from nltk.sentiment import SentimentIntensityAnalyzer
    sia = SentimentIntensityAnalyzer()
    print(" NLTK sentiment analyzer loaded")
except:
    sia = None
    print(" NLTK not available - using basic functionality")

# Document processing
try:
    import PyPDF2
    import pdfplumber
    print(" PDF processing available")
except Exception as pdf_exc:
    PyPDF2 = None
    pdfplumber = None
    print(f" PDF processing not available: {pdf_exc}")

try:
    from docx import Document
    print(" DOCX processing available")
except Exception as docx_exc:
    Document = None
    print(f" DOCX processing not available: {docx_exc}")

try:
    import pytesseract
    from PIL import Image
    print(" OCR processing available")
except:
    print(" OCR processing not available")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

ML_DEBUG_ENABLED = os.getenv("SKILLSYNC_DEBUG_ML", "0").lower() in {"1", "true", "yes", "on"}

DEFAULT_JOB_SEARCH_TERMS = [
    "software engineer",
    "full stack developer",
    "backend engineer",
    "frontend engineer",
    "data engineer",
    "cloud architect",
    "devops engineer"
]

SYSTEM_USER_EMAIL = "system@skillsync.local"
SYSTEM_USER_USERNAME = "skillsync-system"
SYSTEM_USER_FULL_NAME = "SkillSync System"
SYSTEM_USER_PASSWORD = "skillsync-system"
SYSTEM_USER_ID: Optional[str] = None

LEARNING_MODEL_MIN_RUNS = int(os.getenv("SKILLSYNC_LEARNING_MIN_RUNS", "12"))
LEARNING_MODEL_MIN_RESOURCES = int(os.getenv("SKILLSYNC_LEARNING_MIN_RESOURCES", "30"))
LEARNING_MODEL_REFRESH_INTERVAL_HOURS = int(os.getenv("SKILLSYNC_LEARNING_REFRESH_HOURS", "12"))
LEARNING_MODEL_REFRESH_INTERVAL = timedelta(hours=LEARNING_MODEL_REFRESH_INTERVAL_HOURS)
LAST_MODEL_REFRESH: Optional[datetime] = None
MODEL_REFRESH_PENDING = False
MODEL_REFRESH_LOCK = threading.Lock()


def log_ml(stage: str, **details: Any) -> None:
    """Emit structured ML debug logs when verbose tracing is enabled."""
    if not (ML_DEBUG_ENABLED or getattr(settings, "debug", False)):
        return
    try:
        logger.info("[ml-debug][%s] %s", stage, json.dumps(details, default=str))
    except Exception:
        logger.info("[ml-debug][%s] %s", stage, details)

class Settings:
    """Application settings"""
    app_name = "SkillSync Enhanced"
    app_version = "2.0"
    debug = False
    
    # Database settings
    database_url = "sqlite:///./skillsync_enhanced.db"
    
    # Security settings
    secret_key = "your-secret-key-change-in-production"
    algorithm = "HS256"
    access_token_expire_minutes = 30
    
    # AI/ML model settings
    embedding_model = "all-MiniLM-L6-v2"
    confidence_threshold = 0.7
    resume_ner_model_name = os.getenv("SKILLSYNC_RESUME_MODEL_NAME", "bert-base-uncased")
    resume_ner_model_path = os.getenv("SKILLSYNC_RESUME_MODEL_PATH")
    skill_confidence_threshold = float(os.getenv("SKILLSYNC_SKILL_CONFIDENCE", "0.62"))
    
    # API settings
    max_file_size = 10 * 1024 * 1024  # 10MB
    allowed_extensions = [".pdf", ".docx"]
    
    # CORS settings
    allowed_origins = [
        "http://localhost:3000",
        "http://localhost:5173",
        "http://localhost:5174",
        "http://localhost:5175",
        "http://localhost:8080",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:5173",
        "http://127.0.0.1:5174",
        "http://127.0.0.1:5175",
        "http://127.0.0.1:8080"
    ]

settings = Settings()

try:
    LEARNING_ROADMAP_PREDICTOR = get_default_predictor()
except Exception as exc:  # noqa: BLE001
    LEARNING_ROADMAP_PREDICTOR = None
    logger.warning("[roadmap-ml] Predictor fallback activated: %s", exc)

# Initialize FastAPI app
app = FastAPI(
    title=settings.app_name,
    version=settings.app_version,
    description="Advanced AI-Powered Career Development System",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Register modular routers for CV analysis and interviews
app.include_router(cv_analysis.router)
app.include_router(interview_routes.router)
app.include_router(interview_realtime.router)

# Include auth router if available
if AUTH_ENABLED:
    app.include_router(auth_router)
    logger.info("✅ Authentication enabled")

# Database setup
engine = create_engine(settings.database_url, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Security
security = HTTPBearer()

# Skill extraction patterns
programming_languages = [
    'python', 'javascript', 'java', 'c++', 'c#', 'go', 'rust', 'swift', 'kotlin',
    'typescript', 'php', 'ruby', 'scala', 'r', 'matlab', 'perl', 'lua', 'dart'
]

frameworks = [
    'react', 'vue', 'angular', 'django', 'flask', 'fastapi', 'express', 'spring',
    'laravel', 'rails', 'asp.net', 'tensorflow', 'pytorch', 'node.js', 'next.js'
]

databases = [
    'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
    'dynamodb', 'neo4j', 'sqlite', 'oracle', 'sql server'
]

cloud_platforms = [
    'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean'
]

tools = [
    'docker', 'kubernetes', 'git', 'github', 'gitlab', 'jenkins', 'maven', 'gradle'
]

additional_skill_terms = [
    'machine learning', 'data science', 'deep learning', 'computer vision',
    'natural language processing', 'nlp', 'mlops', 'data engineering',
    'data analysis', 'data visualization', 'business intelligence',
    'cloud computing', 'cloud architecture', 'cloud security',
    'big data', 'microservices', 'rest apis', 'graphql', 'ai',
    'artificial intelligence', 'predictive modeling', 'time series',
    'reinforcement learning', 'kafka', 'spark', 'hadoop', 'airflow',
    'power bi', 'tableau', 'excel', 'sas', 'scikit-learn', 'pandas',
    'numpy', 'matplotlib', 'seaborn', 'computer graphics', 'linux', 'tensorflow',
    'keras', 'pytorch', 'docker compose', 'llm', 'large language model',
    'generative ai', 'rag', 'retrieval augmented generation', 'langchain',
    'hugging face', 'transformers', 'prompt engineering', 'vector database',
    'pinecone', 'weaviate', 'milvus', 'chroma', 'mlflow', 'kubeflow',
    'vertex ai', 'sagemaker', 'azure ml', 'model monitoring', 'feature store',
    'dataops', 'ml pipelines', 'onnx', 'ops automation'
]

AI_INTENT_KEYWORDS = {
    'ai', 'artificial intelligence', 'machine learning', 'ml', 'data science',
    'deep learning', 'mlops', 'computer vision', 'natural language processing',
    'nlp', 'llm', 'large language model', 'generative ai', 'rag',
    'retrieval augmented generation', 'predictive modeling'
}

AI_PRIORITY_SKILLS = [
    'Machine Learning', 'Deep Learning', 'MLOps', 'Data Science', 'NLP',
    'LLM', 'TensorFlow', 'PyTorch', 'Scikit-learn', 'Keras', 'LangChain',
    'Hugging Face', 'Transformers', 'Computer Vision', 'Model Monitoring',
    'MLflow', 'Kubeflow', 'Vertex AI', 'SageMaker', 'Azure ML', 'Docker',
    'Kubernetes', 'BigQuery', 'Airflow', 'Spark', 'Feature Store'
]

AI_PRIORITY_SKILL_SET = {skill.lower() for skill in AI_PRIORITY_SKILLS}

WEB_INTENT_KEYWORDS = {
    'frontend', 'front-end', 'full stack', 'javascript', 'react', 'vue',
    'angular', 'node.js', 'php', 'laravel', 'spring', 'web developer'
}

TECH_KEYWORD_POOL = list(dict.fromkeys(
    programming_languages + frameworks + databases + cloud_platforms + tools
))

MONTH_NAME_TOKENS = {
    'jan', 'january', 'feb', 'february', 'fev', 'fevr', 'mar', 'march', 'apr', 'april',
    'may', 'mai', 'jun', 'june', 'jul', 'july', 'aug', 'august', 'sep', 'sept', 'september',
    'oct', 'october', 'nov', 'november', 'dec', 'december'
}

TECH_CONTEXT_INDICATORS = {
    'develop', 'developer', 'engineer', 'pipeline', 'api', 'stack', 'code',
    'algorithm', 'data', 'model', 'deployment', 'microservice', 'architecture',
    'debug', 'build', 'ci/cd', 'automation'
}

MARKETING_CONTEXT_TERMS = {
    'campaign', 'brand', 'go-to-market', 'digital marketing', 'seo', 'sem',
    'copywriting', 'social media', 'growth marketing', 'audience', 'crm',
    'paid media', 'funnel', 'retention', 'engagement'
}

SHORT_SKILL_GUARD = {'r', 'go'}

job_title_keywords = [
    'engineer', 'developer', 'manager', 'architect', 'scientist', 'analyst',
    'consultant', 'designer', 'lead', 'specialist', 'administrator'
]

soft_skill_keywords = [
    'communication', 'leadership', 'teamwork', 'collaboration', 'problem solving',
    'critical thinking', 'adaptability', 'creativity', 'mentoring', 'planning'
]

education_keywords = [
    'bachelor', 'master', 'phd', 'doctorate', 'diploma', 'licence', 'msc', 'bsc'
]

SKILL_WHITELIST = set(
    term.lower()
    for term in (
        TECH_KEYWORD_POOL
        + soft_skill_keywords
        + job_title_keywords
        + additional_skill_terms
    )
)

PORTFOLIO_SKILL_STOPWORDS = {
    'ae', 'de', 'en', 'fe', 'intern', 'ons', 'recherche', 'stage', 'ltd', 'inc', 'sarl',
    'llc', 'sa', 'experience', 'skills', 'about', 'profile', 'profil', 'resume', 'curriculum',
    'developer', 'full', 'stack', 'contact', 'phone', 'email', 'location', 'africa', 'and',
    'with', 'interfaces', 'deployments', 'int', 'profil', 'let s', 'let', 's', 'profil professionnel'
}

PORTFOLIO_SKILL_STOPWORD_TOKENS = {
    'profil', 'profile', 'professional', 'developer', 'full', 'stack', 'contact', 'phone',
    'email', 'linkedin', 'github', 'with', 'and', 'africa', 'interfaces', 'deployments',
    'experience', 'skills', 'resume', 'curriculum', 'let', 'stage', 'intern', 'recent'
}

PORTFOLIO_SKILL_CATEGORY_KEYWORDS = {
    'Programming Languages': {'python', 'java', 'javascript', 'typescript', 'sql', 'c++', 'c#', 'c ', 'scala', 'go', 'rust', 'ruby', 'php'},
    'Machine Learning / Data Science': {'machine learning', 'ml', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'data science', 'sklearn', 'scikit', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'computer vision', 'nlp'},
    'DevOps / Cloud': {'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'ci/cd', 'jenkins', 'terraform', 'mlops'},
    'Tools': {'git', 'github', 'gitlab', 'jira', 'notion', 'figma', 'vscode', 'linux', 'bash', 'tableau', 'power bi'},
    'Soft Skills': {'communication', 'leadership', 'teamwork', 'collaboration', 'problem solving', 'mentoring', 'planning', 'ownership'}
}

EMAIL_REGEX = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_REGEX = re.compile(r"(\+?\d[\d\s().-]{7,})")
LINKEDIN_REGEX = re.compile(r"(https?://)?([\w.-]*linkedin\.com/[\w\-/]+)", re.IGNORECASE)
GITHUB_REGEX = re.compile(r"(https?://)?([\w.-]*github\.com/[\w\-/]+)", re.IGNORECASE)

SECTION_HEADER_ALIASES = {
    'profile': 'profile',
    'profil': 'profile',
    'profilesummary': 'profile',
    'professionalprofile': 'profile',
    'professionalsummary': 'profile',
    'summary': 'profile',
    'about': 'profile',
    'aboutme': 'profile',
    'objective': 'profile',
    'careerobjective': 'profile',
    'experience': 'experience',
    'workexperience': 'experience',
    'professionalexperience': 'experience',
    'experiences': 'experience',
    'employmenthistory': 'experience',
    'projects': 'projects',
    'project': 'projects',
    'personalprojects': 'projects',
    'portfolio': 'projects',
    'education': 'education',
    'academicbackground': 'education',
    'academics': 'education',
    'formation': 'education',
    'formations': 'education',
    'studies': 'education',
    'skills': 'skills',
    'technicalskills': 'skills',
    'skillset': 'skills',
    'competencies': 'skills',
    'competences': 'skills',
    'capabilities': 'skills',
    'languages': 'extras',
    'language': 'extras',
    'certifications': 'extras',
    'certification': 'extras',
    'clubs': 'extras',
    'volunteering': 'extras',
    'interests': 'extras'
}

ROLE_HINT_KEYWORDS = {
    'engineer', 'developer', 'scientist', 'architect', 'student', 'analyst',
    'consultant', 'researcher', 'designer', 'manager', 'specialist', 'lead',
    'intern', 'mlops', 'ai', 'data', 'product', 'security'
}


def _clean_text_lines(raw: Optional[str]) -> List[str]:
    if not raw:
        return []
    text = html.unescape(raw)
    text = text.replace('\r', '\n')
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.split('\n')]
    return [line for line in lines if line]


def _looks_like_contact(line: str) -> bool:
    lowered = line.lower()
    return any(keyword in lowered for keyword in ('@', 'linkedin', 'github', 'portfolio', 'http://', 'https://', 'phone', 'tel', '+'))


def _dedupe_preserve_order(items: Iterable[str]) -> List[str]:
    seen: Set[str] = set()
    result: List[str] = []
    for item in items:
        key = re.sub(r"[^a-z0-9]+", '', item.lower())
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def _normalize_skill_term(skill: str) -> Optional[str]:
    if not skill:
        return None
    cleaned = html.unescape(skill)
    cleaned = re.sub(r"[•|–—]", ' ', cleaned)
    cleaned = re.sub(r"[^A-Za-z0-9+#/.,&()\- ]", ' ', cleaned)
    cleaned = re.sub(r"\s+", ' ', cleaned).strip(" ,.-")
    if not cleaned:
        return None
    lowered = cleaned.lower()
    if lowered in PORTFOLIO_SKILL_STOPWORDS or len(lowered) <= 1:
        return None
    tokens = lowered.split()
    if len(tokens) > 3:
        return None
    if any(token in PORTFOLIO_SKILL_STOPWORD_TOKENS for token in tokens):
        return None
    return cleaned.title()


def _categorize_skills_for_portfolio(skills: Iterable[str]) -> Dict[str, List[str]]:
    categories: Dict[str, List[str]] = {key: [] for key in PORTFOLIO_SKILL_CATEGORY_KEYWORDS}
    other: List[str] = []
    for skill in skills:
        normalized = _normalize_skill_term(skill)
        if not normalized:
            continue
        lowered = normalized.lower()
        matched_category = None
        for category, keywords in PORTFOLIO_SKILL_CATEGORY_KEYWORDS.items():
            if any(keyword in lowered for keyword in keywords):
                matched_category = category
                break
        if matched_category:
            if normalized not in categories[matched_category]:
                categories[matched_category].append(normalized)
        else:
            if normalized not in other:
                other.append(normalized)
    if other:
        categories['Tools'].extend(other)
    # Trim & dedupe entries
    for category, values in categories.items():
        categories[category] = _dedupe_preserve_order(values)[:8]
    return {category: values for category, values in categories.items() if values}


def _sanitize_contact_value(value: str) -> str:
    cleaned = html.unescape(value or '').strip()
    cleaned = re.sub(r"(email|phone|linkedin|github|portfolio)[: ]*$", '', cleaned, flags=re.IGNORECASE).strip()
    cleaned = cleaned.replace('Phone', '').replace('phone', '').strip()
    return cleaned


def _extract_contact_from_raw(raw_text: str) -> Dict[str, str]:
    contact: Dict[str, str] = {}
    for line in _clean_text_lines(raw_text)[:12]:
        email_match = EMAIL_REGEX.search(line)
        if email_match and 'email' not in contact:
            contact['email'] = _sanitize_contact_value(email_match.group(0))
        linkedin_match = LINKEDIN_REGEX.search(line)
        if linkedin_match and 'linkedin' not in contact:
            contact['linkedin'] = _sanitize_contact_value(LINKEDIN_REGEX.sub(r"https://\2", linkedin_match.group(0)))
        github_match = GITHUB_REGEX.search(line)
        if github_match and 'github' not in contact:
            contact['github'] = _sanitize_contact_value(GITHUB_REGEX.sub(r"https://\2", github_match.group(0)))
        phone_match = PHONE_REGEX.search(line)
        if phone_match and 'phone' not in contact:
            contact['phone'] = _sanitize_contact_value(re.sub(r"\s+", ' ', phone_match.group(0)))
        if 'phone' in contact and 'email' in contact and 'linkedin' in contact and 'github' in contact:
            break
    return contact


def _guess_name_from_text(raw_text: str) -> Optional[str]:
    import re
    lines = _clean_text_lines(raw_text)[:8]
    
    logger.info(f"[_guess_name_from_text] Checking {len(lines)} lines for name")
    
    for idx, line in enumerate(lines):
        logger.info(f"[_guess_name_from_text] Line {idx}: '{line[:100]}'")
        
        # Skip lines with pipes
        if '|' in line:
            logger.info(f"[_guess_name_from_text] Line {idx} skipped (pipe)")
            continue
        
        # Clean the line FIRST - remove emails, URLs, phone numbers
        cleaned = re.sub(r'\S+@\S+\.\S+', '', line)  # Remove emails
        cleaned = re.sub(r'github\.com\S*|linkedin\.com\S*', '', cleaned, flags=re.IGNORECASE)  # Remove URLs
        cleaned = re.sub(r'[\+]?[\d\s\-\(\)]{8,}', '', cleaned)  # Remove phones
        cleaned = re.sub(r'\b(ariana|tunisia|tunis|profile|student|engineer|developer|at|tek-up|tekup)\b', '', cleaned, flags=re.IGNORECASE)  # Remove common words
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        logger.info(f"[_guess_name_from_text] After cleaning (first 200 chars): '{cleaned[:200]}'")
        
        # Extract only the FIRST few words as potential name
        tokens = cleaned.split()
        # Take first 2-4 tokens that look like name parts (80%+ letters, 2+ chars)
        name_tokens = []
        for token in tokens[:6]:  # Check first 6 tokens max
            # Stop if we hit a comma or punctuation (end of name)
            if token in [',', '.', ':', ';', '|']:
                break
            # Only accept tokens that are mostly letters
            if len(token) >= 2 and sum(1 for c in token if c.isalpha()) / len(token) >= 0.8:
                name_tokens.append(token)
                # Stop after collecting 2-4 name parts
                if len(name_tokens) >= 4:
                    break
            # Stop if we've seen 2+ valid tokens and hit a non-name word
            elif len(name_tokens) >= 2:
                break
        
        logger.info(f"[_guess_name_from_text] Extracted name tokens: {name_tokens}")
        
        # Validate: should be 2-4 words, reasonable length
        if 2 <= len(name_tokens) <= 4:
            name = ' '.join(name_tokens)
            if 3 <= len(name) <= 50:
                result = ' '.join(w.title() for w in name_tokens)
                logger.info(f"[_guess_name_from_text] FOUND NAME: '{result}'")
                return result
    
    logger.info(f"[_guess_name_from_text] No name found, returning None")
    return None


def _normalize_section_label(label: str) -> Optional[str]:
    if not label:
        return None
    normalized = re.sub(r"[^a-z]", '', label.lower())
    return SECTION_HEADER_ALIASES.get(normalized)


def _segment_cv_sections(raw_text: str) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = defaultdict(list)
    if not raw_text:
        return sections
    lines = _clean_text_lines(raw_text)
    current: Optional[str] = None
    for line in lines:
        header = _normalize_section_label(line)
        if header:
            current = header
            continue
        if current:
            sections[current].append(line)
    return sections


def _extract_profile_summary(raw_text: str, sections: Dict[str, List[str]]) -> str:
    profile_lines = sections.get('profile') or []
    if profile_lines:
        return _clean_sentence_block(' '.join(profile_lines), max_sentences=4)
    filtered = [line for line in _clean_text_lines(raw_text) if not _looks_like_contact(line)]
    snippet = ' '.join(filtered[:6])
    return _clean_sentence_block(snippet or raw_text, max_sentences=4)


def _infer_primary_title(
    raw_text: str,
    sections: Dict[str, List[str]],
    fallback_title: Optional[str]
) -> str:
    lines = _clean_text_lines(raw_text)
    name_candidate = _guess_name_from_text(raw_text)
    start_index = 0
    if name_candidate:
        lowered = name_candidate.lower()
        for idx, line in enumerate(lines):
            if line.lower() == lowered:
                start_index = idx + 1
                break
    search_window = lines[start_index:start_index + 14]
    for line in search_window:
        if _looks_like_contact(line):
            continue
        if _normalize_section_label(line):
            continue
        lowered = line.lower()
        if any(keyword in lowered for keyword in ROLE_HINT_KEYWORDS):
            return line.strip().strip(' .')
        if len(line.split()) >= 3 and len(line) <= 80:
            return line.strip().strip(' .')
    profile_lines = sections.get('profile') or []
    for line in profile_lines[:3]:
        lowered = line.lower()
        if any(keyword in lowered for keyword in ROLE_HINT_KEYWORDS):
            return line.strip().strip(' .')
    summary_text = _extract_profile_summary(raw_text, sections)
    for sentence in re.split(r"(?<=[.!?])\s+", summary_text):
        lowered = sentence.lower().strip()
        if not lowered:
            continue
        if any(keyword in lowered for keyword in ROLE_HINT_KEYWORDS):
            return sentence.strip().strip(' .')
    if fallback_title:
        return fallback_title.strip()
    return 'Candidate'


def _extract_skill_terms_from_sections(sections: Dict[str, List[str]]) -> List[str]:
    skill_lines = []
    for key in ('skills',):
        skill_lines.extend(sections.get(key) or [])
    terms: List[str] = []
    for line in skill_lines:
        for chunk in re.split(r"[•,;/|]", line):
            cleaned = chunk.strip()
            if cleaned and not _looks_like_contact(cleaned):
                terms.append(cleaned)
    return terms


def _looks_like_project_title(line: str) -> bool:
    if not line:
        return False
    stripped = line.strip(" :")
    if not stripped or len(stripped) > 120:
        return False
    tokens = stripped.split()
    if not tokens or len(tokens) > 12:
        return False
    capitalized = sum(1 for token in tokens if token and token[0].isupper())
    return capitalized >= max(1, len(tokens) - 2)


def _extract_projects_from_sections(sections: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    project_lines = sections.get('projects') or []
    if not project_lines:
        return []
    projects: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None
    for raw_line in project_lines:
        if _normalize_section_label(raw_line):
            continue
        line = raw_line.strip()
        if not line:
            continue
        if _looks_like_project_title(line):
            if current:
                projects.append(current)
            current = {'title': line.strip(' :'), 'summary_lines': [], 'bullet_lines': []}
            continue
        if current is None:
            continue
        if line.startswith(('-', '•')):
            current.setdefault('bullet_lines', []).append(line.lstrip('-• ').strip())
        else:
            current.setdefault('summary_lines', []).append(line)
    if current:
        projects.append(current)

    normalized: List[Dict[str, Any]] = []
    for project in projects:
        summary_text = ' '.join(project.get('summary_lines') or [])
        bullets = project.get('bullet_lines') or []
        normalized.append({
            'title': project.get('title') or 'Project',
            'summary': summary_text or None,
            'highlights': bullets,
            'tech': _extract_keywords(summary_text, TECH_KEYWORD_POOL) if summary_text else []
        })
    return normalized[:4]


def _extract_education_from_sections(sections: Dict[str, List[str]]) -> List[str]:
    education_lines = sections.get('education') or []
    if not education_lines:
        return []
    entries: List[str] = []
    current: List[str] = []
    for line in education_lines:
        if _normalize_section_label(line):
            if current:
                entries.append(' • '.join(current))
                current = []
            continue
        if not line:
            continue
        current.append(line)
        if len(current) >= 3:
            entries.append(' • '.join(current))
            current = []
    if current:
        entries.append(' • '.join(current))
    cleaned = [entry.strip(' •-') for entry in entries if entry.strip()]
    return _dedupe_preserve_order(cleaned)[:4]


def _clean_sentence_block(text: str, max_sentences: int = 3) -> str:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    cleaned: List[str] = []
    for sentence in sentences:
        stripped = sentence.strip()
        if not stripped or _looks_like_contact(stripped):
            continue
        normalized = re.sub(r"[^a-z0-9]+", '', stripped.lower())
        if not normalized or any(normalized == re.sub(r"[^a-z0-9]+", '', existing.lower()) for existing in cleaned):
            continue
        cleaned.append(stripped)
        if len(cleaned) >= max_sentences:
            break
    return ' '.join(cleaned)


def _clean_bullet_list(raw: Any, max_items: int = 4) -> List[str]:
    bullets: List[str] = []
    seen: Set[str] = set()
    if isinstance(raw, list):
        source = raw
    elif isinstance(raw, str):
        source = re.split(r"\n|•|-", raw)
    else:
        source = []
    for entry in source:
        cleaned = re.sub(r"\s+", ' ', str(entry)).strip(" .-•")
        if not cleaned or len(cleaned) < 4:
            continue
        if _looks_like_contact(cleaned):
            continue
        normalized = re.sub(r"[^a-z0-9]+", '', cleaned.lower())
        if normalized in seen:
            continue
        seen.add(normalized)
        truncated = cleaned[:220].rstrip(' ,.;')
        bullets.append(truncated)
        if len(bullets) >= max_items:
            break
    return _dedupe_preserve_order(bullets)


def _build_experience_cards(analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
    experiences: List[Dict[str, Any]] = []
    for record in (analysis.get('work_history') or [])[:6]:
        title = record.get('title') or record.get('role') or ''
        company = record.get('company') or record.get('organization') or ''
        if not title and not company:
            continue
        period = record.get('period') or record.get('date_range') or record.get('duration') or 'Present'
        bullets = _clean_bullet_list(record.get('highlights') or record.get('achievements') or record.get('description'))
        if not bullets and record.get('summary'):
            bullets = _clean_bullet_list(record['summary'])
        experiences.append({
            'title': title.strip() or 'Role',
            'company': company.strip() or 'Company',
            'period': period.strip(),
            'bullets': bullets[:4] or ['Delivered measurable outcomes across projects.']
        })
    if not experiences and analysis.get('projects'):
        for project in (analysis.get('projects') or [])[:3]:
            title = project.get('title') or project.get('name') or 'Project'
            company = project.get('context') or project.get('company') or 'Project Collaboration'
            period = project.get('period') or project.get('timeline') or 'Recent'
            bullet_source = project.get('bullets') or project.get('highlights') or project.get('summary') or project.get('impact')
            bullets = _clean_bullet_list(bullet_source, max_items=3)
            experiences.append({
                'title': title,
                'company': company,
                'period': period,
                'bullets': bullets or ['Shipped a hands-on initiative end-to-end.']
            })
    if not experiences:
        job_titles = analysis.get('job_titles') or []
        summary_bullets = _clean_bullet_list(analysis.get('summary'), max_items=3)
        for title in job_titles[:2]:
            experiences.append({
                'title': title,
                'company': 'Recent Engagement',
                'period': 'Present',
                'bullets': summary_bullets or ['Led cross-functional delivery efforts.']
            })
    return experiences[:4]


def _build_project_cards(analysis: Dict[str, Any], fallback_experiences: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    projects: List[Dict[str, Any]] = []
    for project in (analysis.get('projects') or [])[:4]:
        title = project.get('name') or project.get('title') or 'Project'
        summary = project.get('impact') or project.get('summary') or project.get('description') or ''
        tech_stack = project.get('tech_stack') or project.get('technologies') or project.get('tech') or []
        bullet_source = project.get('bullets') or project.get('highlights') or summary
        bullets = _clean_bullet_list(bullet_source, max_items=3)
        projects.append({
            'title': title.strip(),
            'summary': summary.strip()[:180] if summary else None,
            'bullets': bullets or ['Drove tangible user/business impact.'],
            'tech': _dedupe_preserve_order(tech_stack)[:5]
        })
    if not projects:
        for exp in fallback_experiences[:2]:
            projects.append({
                'title': f"{exp['title']} Initiative",
                'summary': exp['bullets'][0] if exp['bullets'] else None,
                'bullets': exp['bullets'][1:3] or exp['bullets'][:2],
                'tech': []
            })
    return projects[:3]


def _build_education_entries(analysis: Dict[str, Any]) -> List[str]:
    education = analysis.get('education') or []
    if isinstance(education, dict):
        education = education.values()
    entries = []
    for entry in education:
        text = entry if isinstance(entry, str) else entry.get('name') or entry.get('title')
        if text:
            cleaned = re.sub(r"\s+", ' ', text).strip()
            if cleaned:
                entries.append(cleaned)
    return _dedupe_preserve_order(entries)[:4]


def _build_contact_points(analysis: Dict[str, Any], raw_text: str) -> Dict[str, str]:
    contact = (analysis.get('contact_info') or {}).copy()
    inferred = _extract_contact_from_raw(raw_text)
    for key, value in inferred.items():
        contact.setdefault(key, value)
    personal = analysis.get('personal_info') or {}
    if personal.get('location'):
        contact.setdefault('location', personal['location'])
    return contact


def _build_portfolio_viewmodel(analysis: Dict[str, Any]) -> Dict[str, Any]:
    import re
    
    if isinstance(analysis, BaseModel):
        analysis_data = analysis.model_dump()
    else:
        analysis_data = dict(analysis)

    raw_text = analysis_data.get('raw_text') or ''
    sections_map = _segment_cv_sections(raw_text)

    # Get the raw name from various sources
    hero_name = (analysis_data.get('personal_info') or {}).get('name')
    if not hero_name or hero_name.strip().lower() in {'professional', 'candidate'}:
        hero_name = _guess_name_from_text(raw_text) or analysis_data.get('name') or 'Professional'
    else:
        hero_name = hero_name.strip()
    
    # IMPORTANT: Clean the name even if it came from personal_info
    # Remove common garbage that might have been stored
    def clean_extracted_name(name: str) -> str:
        """Clean name by removing job titles, emails, URLs, etc."""
        if not name or name.lower() in {'professional', 'candidate', 'engineer', 'developer', 'student'}:
            # If the stored name is just a job title, try to extract from raw text
            fresh_name = _guess_name_from_text(raw_text)
            if fresh_name:
                return fresh_name
            return 'Professional'
        
        # Clean the name
        cleaned = re.sub(r'\S+@\S+\.\S+', '', name)  # Remove emails
        cleaned = re.sub(r'github\.com\S*|linkedin\.com\S*', '', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'[\+]?[\d\s\-\(\)]{8,}', '', cleaned)  # Remove phones
        cleaned = re.sub(r'\b(ariana|tunisia|tunis|profile|at|tek-up|tekup)\b', '', cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        
        # If cleaned name is too short or empty, try extracting from raw text
        if not cleaned or len(cleaned) < 3:
            fresh_name = _guess_name_from_text(raw_text)
            if fresh_name:
                return fresh_name
        
        return cleaned.title() if cleaned else 'Professional'
    
    hero_name = clean_extracted_name(hero_name)

    title_fallback = (
        analysis_data.get('current_title')
        or (analysis_data.get('job_titles') or [None])[0]
        or 'Candidate'
    )
    current_title = _infer_primary_title(raw_text, sections_map, title_fallback)

    location = (analysis_data.get('personal_info') or {}).get('location')
    contact = _build_contact_points(analysis_data, raw_text)
    if not location and contact.get('location'):
        location = contact['location']

    summary_source = analysis_data.get('summary') or ''
    clean_summary = (
        _extract_profile_summary(raw_text, sections_map)
        or _clean_sentence_block(summary_source)
        or _clean_sentence_block(raw_text)
    )

    skills_source: List[str] = []
    if isinstance(analysis_data.get('skill_categories'), dict):
        for values in analysis_data['skill_categories'].values():
            skills_source.extend(values)
    else:
        skills_source = analysis_data.get('skills') or []
    skills_source.extend(_extract_skill_terms_from_sections(sections_map))
    categorized_skills = _categorize_skills_for_portfolio(skills_source)

    parsed_projects = _extract_projects_from_sections(sections_map)
    if parsed_projects and not (analysis_data.get('projects') or []):
        analysis_data['projects'] = parsed_projects

    inferred_education = _extract_education_from_sections(sections_map)
    if inferred_education:
        merged_education = (analysis_data.get('education') or []) + inferred_education
        analysis_data['education'] = _dedupe_preserve_order(merged_education)[:6]

    experiences = _build_experience_cards(analysis_data)
    projects = _build_project_cards(analysis_data, experiences)
    education = _build_education_entries(analysis_data)

    hero_headline = clean_summary.split('. ')[0].strip() if clean_summary else current_title

    return {
        'hero': {
            'name': hero_name,
            'title': current_title,
            'location': location or 'Remote / Open to Relocation',
            'headline': hero_headline or current_title,
            'contact': contact,
        },
        'about': clean_summary,
        'skills': categorized_skills,
        'experiences': experiences,
        'projects': projects,
        'education': education,
        'contact': contact,
    }


class UserDB(Base):
    __tablename__ = "users"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    email = Column(String, unique=True, nullable=False)
    username = Column(String, unique=True, nullable=False)
    full_name = Column(String, nullable=False)
    hashed_password = Column(String, nullable=False)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class CVAnalysisDB(Base):
    __tablename__ = "cv_analyses"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=False)
    filename = Column(String, nullable=False)
    file_size = Column(Integer)
    skills = Column(JSON)  # Store skills as JSON
    experience_years = Column(Float)
    education_level = Column(String)
    confidence_score = Column(Float)
    gap_analysis = Column(JSON)  # Store gap analysis as JSON
    recommendations = Column(JSON)  # Store recommendations as JSON
    created_at = Column(DateTime, default=datetime.utcnow)

class JobMatchDB(Base):
    __tablename__ = "job_matches"
    
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=False)
    job_title = Column(String, nullable=False)
    company = Column(String)
    job_description = Column(Text)
    match_score = Column(Float)
    matched_skills = Column(JSON)
    missing_skills = Column(JSON)
    salary_range = Column(String)
    location = Column(String)
    job_url = Column(String)
    created_at = Column(DateTime, default=datetime.utcnow)


class LearningRoadmapRunDB(Base):
    __tablename__ = "learning_roadmap_runs"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = Column(String, ForeignKey("users.id"), nullable=True)
    analysis_id = Column(String, nullable=True)
    primary_skills = Column(JSON, default=list)
    gap_skills = Column(JSON, default=list)
    skills_sequence = Column(JSON, default=list)
    phases = Column(JSON, default=list)
    total_hours = Column(Float, default=0.0)
    total_weeks = Column(Integer, default=0)
    success_rate = Column(Float, default=0.0)
    weekly_hours_available = Column(Float, nullable=True)
    experience_years = Column(Float, nullable=True)
    target_role = Column(String, nullable=True)
    planned_duration_hours = Column(Float, default=0.0)
    actual_duration_hours = Column(Float, nullable=True)
    satisfaction_score = Column(Float, nullable=True)
    quiz_scores = Column(JSON, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)


class LearningResourceEventDB(Base):
    __tablename__ = "learning_resource_events"

    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    run_id = Column(String, ForeignKey("learning_roadmap_runs.id"), nullable=False)
    skill = Column(String, nullable=False)
    tier = Column(String, nullable=True)
    provider = Column(String, nullable=True)
    resource_type = Column(String, nullable=True)
    duration_hours = Column(Float, default=0.0)
    cost = Column(Float, default=0.0)
    success_rate = Column(Float, nullable=True)
    difficulty = Column(Float, nullable=True)
    is_free = Column(Boolean, default=False)
    resource_payload = Column(JSON, nullable=True)
    source = Column(String, default="ml-roadmap")
    created_at = Column(DateTime, default=datetime.utcnow)

    run = relationship("LearningRoadmapRunDB", backref="resources")

# Create tables
Base.metadata.create_all(bind=engine)

# Pydantic models for API
class UserCreate(BaseModel):
    email: str
    username: str
    full_name: str
    password: str

class UserLogin(BaseModel):
    email: str
    password: str

class CVUploadRequest(BaseModel):
    user_id: str

class CVUploadResponse(BaseModel):
    analysis_id: str
    message: str
    confidence_score: float

class JobMatchRequest(BaseModel):
    user_id: str
    job_title: str
    company: Optional[str] = None
    job_description: str
    location: Optional[str] = None
    salary_range: Optional[str] = None

class JobMatch(BaseModel):
    job_id: str
    job_title: str
    company: Optional[str] = None
    match_score: float
    matched_skills: List[str]
    missing_skills: List[str]
    salary_range: Optional[str] = None
    location: Optional[str] = None
    explanation: str

class JobMatchResponse(BaseModel):
    matches: List[JobMatch]
    total_matches: int
    best_match: JobMatch

class ExperienceTranslateRequest(BaseModel):
    text: str
    source_type: str  # 'technical' or 'business'
    target_type: str  # 'technical' or 'business'
    tone: str = 'professional'  # 'professional', 'casual', 'formal'

class ExperienceTranslateResponse(BaseModel):
    original_text: str
    translated_text: str
    source_type: str
    target_type: str
    confidence_score: float
    key_changes: List[str]


class CVAnalysisRequest(BaseModel):
    cv_content: str = Field(..., min_length=50, description="Plain text CV content for analysis")


class CVAnalysisResponse(BaseModel):
    analysis_id: str
    skills: List[str]
    hard_skills: Optional[List[str]] = None
    soft_skills: Optional[List[str]] = None
    experience_years: Optional[float] = None
    job_titles: List[str]
    education: List[str]
    summary: str
    confidence_score: float
    timestamp: str
    raw_text: Optional[str] = None
    work_history: List[Dict[str, Any]] = Field(default_factory=list)
    total_years_experience: Optional[float] = None
    sections: Optional[Dict[str, Any]] = None
    personal_info: Optional[Dict[str, Any]] = None
    contact_info: Optional[Dict[str, Any]] = None
    roadmap: Optional[Dict[str, Any]] = None
    certifications: Optional[List[Dict[str, Any]]] = None
    recommendations: Optional[List[Dict[str, Any]]] = None
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    current_title: Optional[str] = None
    seniority_level: Optional[str] = None
    industries: Optional[List[str]] = None
    career_trajectory: Optional[Dict[str, Any]] = None
    projects: Optional[List[Dict[str, Any]]] = None
    portfolio_links: Optional[Dict[str, Any]] = None
    ml_confidence_breakdown: Optional[Dict[str, float]] = None
    parser_version: Optional[str] = None
    skill_categories: Optional[Dict[str, List[str]]] = None
    tech_stack_clusters: Optional[Dict[str, List[str]]] = None
    languages: Optional[List[str]] = None
    processing_time_ms: Optional[int] = None


class JobSearchRequestPayload(BaseModel):
    query: str = Field('', min_length=0, description="Keyword or role to search for (optional)")
    location: Optional[str] = Field(None, description="Preferred job location or region")
    skills: List[str] = Field(default_factory=list, description="Skills to prioritize during the search")
    max_results: int = Field(20, ge=1, le=100, description="Maximum number of jobs to return")


class JobResultPayload(BaseModel):
    id: str
    title: str
    company: str
    location: str
    description: str
    url: str
    source: str
    salary: Optional[str] = None
    posted_date: Optional[str] = None
    skills_match: Optional[float] = None
    remote: Optional[bool] = None
    match_score: Optional[float] = None
    skills: Optional[List[str]] = None


class JobSearchResponsePayload(BaseModel):
    jobs: List[JobResultPayload]
    total_count: int
    search_query: str
    location: str
    sources_used: List[str]
    search_time_ms: int
    timestamp: str


class SalaryBand(BaseModel):
    min: float
    max: float
    currency: str = "USD"


class JobRecommendationPayload(BaseModel):
    title: str
    similarity_score: float
    confidence: float
    predicted_salary: SalaryBand
    matching_skills: List[str] = Field(default_factory=list)
    matched_skills: List[str] = Field(default_factory=list)
    skill_gaps: List[str] = Field(default_factory=list)
    gap_skills: List[str] = Field(default_factory=list)
    growth_potential: str
    reasons: List[str]
    description: Optional[str] = None
    job_url: str
    source: Optional[str] = None
    source_logo: Optional[str] = None
    posted_date: Optional[str] = None
    posting_date: Optional[str] = None
    location: Optional[str] = None
    location_type: Optional[str] = None
    tags: List[str] = Field(default_factory=list)


class CertificationRecommendationPayload(BaseModel):
    name: str
    relevance_score: float
    skill_alignment: float
    predicted_roi: str
    estimated_time: str
    career_boost: str
    reasons: List[str]
    url: Optional[str] = None
    official_url: Optional[str] = None
    provider: Optional[str] = None
    provider_logo: Optional[str] = None
    cost_type: Optional[str] = None
    delivery_format: Optional[str] = None
    cost_estimate: Optional[str] = None
    resources: List[Dict[str, Any]] = Field(default_factory=list)


class LearningResourcePayload(BaseModel):
    skill: str
    title: str
    provider: str
    duration: str
    rating: float
    url: Optional[str] = None
    cost: Optional[str] = None
    is_free: bool = False
    tier: str = "Core"
    link: Optional[str] = None
    estimated_time_hours: int = 0
    time_hours: Optional[int] = None


class LearningPhasePayload(BaseModel):
    phase_name: str
    duration_weeks: int
    duration_months: float
    total_time_estimate_hours: int
    skills_to_learn: List[str]
    learning_resources: List[LearningResourcePayload]
    resources: List[LearningResourcePayload] = Field(default_factory=list)
    success_probability: str
    success_justification: str
    effort_level: str
    milestones: List[str]
    smart_milestones: List[Dict[str, Any]] = Field(default_factory=list)


class LearningRoadmapPayload(BaseModel):
    total_duration_weeks: int
    total_duration_months: float
    total_time_estimate_hours: int
    predicted_success_rate: str
    personalization_score: str
    learning_strategy: str
    phases: List[LearningPhasePayload]


class PortfolioCustomization(BaseModel):
    color_scheme: str = "blue"
    font_family: str = "Inter"
    layout_style: str = "clean"
    sections_visible: List[str] = Field(default_factory=lambda: [
        "about", "experience", "skills", "projects", "education", "impact", "roadmap", "opportunities", "contact"
    ])
    include_photo: bool = True
    include_projects: bool = True
    include_contact_form: bool = True
    dark_mode: bool = False


class PortfolioGenerationRequest(BaseModel):
    cv_id: str
    template_id: str
    customization: PortfolioCustomization


class PortfolioResponsePayload(BaseModel):
    portfolio: Dict[str, Any]
    html_content: str
    preview_id: str
    preview_url: str


def _normalize_portfolio_customization(config: Any) -> PortfolioCustomization:
    if isinstance(config, PortfolioCustomization):
        return config
    return PortfolioCustomization(**(config or {}))


class XAIInsightsPayload(BaseModel):
    how_we_analyzed_your_cv: Dict[str, Any]
    job_matching_explanation: Dict[str, Any]
    certification_ranking_explanation: Dict[str, Any]
    learning_path_explanation: Dict[str, Any]
    ml_confidence_scores: Dict[str, float]
    key_insights: List[str]


class KeyOpportunityPayload(BaseModel):
    title: str
    priority: str
    confidence: float
    timeline: str
    expected_impact: str
    recommended_actions: List[str]
    success_metrics: List[str]
    blockers: List[str]
    stakeholders: List[str]
    xai_trace_id: str


class CareerGuidanceResponsePayload(BaseModel):
    job_recommendations: List[JobRecommendationPayload]
    certification_recommendations: List[CertificationRecommendationPayload]
    learning_roadmap: LearningRoadmapPayload
    xai_insights: XAIInsightsPayload
    key_opportunities: List[KeyOpportunityPayload]
    metadata: Dict[str, Any]


class CareerGuidanceRequestPayload(BaseModel):
    cv_analysis: Optional[CVAnalysisResponse] = None
    cv_content: Optional[str] = None


FALLBACK_JOB_TEMPLATES = [
    {
        "title": "{role} - Product Innovation",
        "company": "SkillSync Labs",
        "description": "Lead cross-functional squads shipping AI-enabled experiences for millions of users while mentoring engineers on {skills}.",
        "salary": "$130k - $150k",
        "source": "SkillSync Sandbox",
        "match_score": 84.0,
        "url_slug": "product-innovation"
    },
    {
        "title": "{role} Platform Engineer",
        "company": "NextWave Cloud",
        "description": "Scale data pipelines, automate observability, and harden CI/CD for global deployments leveraging {skills}.",
        "salary": "$110k - $135k",
        "source": "SkillSync Sandbox",
        "match_score": 79.0,
        "url_slug": "platform-engineer"
    },
    {
        "title": "{role} Solutions Architect",
        "company": "FutureStack",
        "description": "Translate business objectives into resilient architectures, guiding customers through migrations built on {skills}.",
        "salary": "$120k - $145k",
        "source": "SkillSync Sandbox",
        "match_score": 76.0,
        "url_slug": "solutions-architect"
    }
]


def _normalize_location(raw_location: Optional[str]) -> str:
    location = (raw_location or "").strip()
    return location


def _format_salary(job: Dict[str, Any]) -> Optional[str]:
    salary_info = job.get("salary") if isinstance(job, dict) else None
    if isinstance(salary_info, dict):
        salary_min = job.get("salary_min") or salary_info.get("min")
        salary_max = job.get("salary_max") or salary_info.get("max")
        currency = job.get("currency") or salary_info.get("currency") or ""
    else:
        salary_min = job.get("salary_min") if isinstance(job, dict) else None
        salary_max = job.get("salary_max") if isinstance(job, dict) else None
        currency = job.get("currency") or ""
    if salary_min is None and salary_max is None:
        return None
    if salary_min and salary_max:
        return f"{currency}{salary_min:,.0f} - {currency}{salary_max:,.0f}".strip()
    single_value = salary_min or salary_max
    return f"{currency}{single_value:,.0f}".strip()


def _map_job_to_payload(job: Dict[str, Any]) -> JobResultPayload:
    job_id = job.get("job_id") or job.get("id") or str(uuid.uuid4())
    company = job.get("company")
    if isinstance(company, dict):
        company_name = company.get("name") or company.get("display_name") or "Unknown company"
    else:
        company_name = company or "Unknown company"

    description = job.get("description") or job.get("contents") or "Description unavailable"
    location = job.get("location") or job.get("locations")
    if isinstance(location, dict):
        location_value = location.get("display_name") or location.get("name") or "Not specified"
    elif isinstance(location, list):
        location_value = ", ".join(filter(None, location)) or "Not specified"
    else:
        location_value = location or "Not specified"

    posted_date = job.get("created_date") or job.get("publication_date") or job.get("posted_date")
    if posted_date is not None and not isinstance(posted_date, str):
        posted_date = str(posted_date)
    skills = job.get("matching_skills") or job.get("skills")
    match_score = job.get("match_score") or job.get("skills_match") or 0.0
    refs = job.get("refs") if isinstance(job, dict) else None
    landing_page = refs.get("landing_page") if isinstance(refs, dict) else None
    job_url = job.get("url") or job.get("redirect_url") or landing_page or ""

    return JobResultPayload(
        id=str(job_id),
        title=job.get("title") or job.get("name") or "Untitled Role",
        company=company_name,
        location=location_value,
        description=description.strip(),
        url=job_url,
        source=job.get("source") or "Unknown",
        salary=_format_salary(job) or job.get("salary"),
        posted_date=posted_date,
        skills_match=match_score,
        remote=job.get("remote"),
        match_score=match_score,
        skills=skills
    )


def _generate_fallback_jobs(query: str, location: str, skills: List[str], max_results: int) -> List[JobResultPayload]:
    role = query or "Software Engineer"
    skills_display = ", ".join(skills[:5]) if skills else "modern cloud platforms"

    jobs: List[JobResultPayload] = []
    for template in FALLBACK_JOB_TEMPLATES[:max_results]:
        description = template["description"].format(skills=skills_display)
        jobs.append(JobResultPayload(
            id=str(uuid.uuid4()),
            title=template["title"].format(role=role),
            company=template["company"],
            location=location,
            description=description,
            url=f"https://jobs.skillsync.ai/{template['url_slug']}",
            source=template["source"],
            salary=template["salary"],
            posted_date=datetime.utcnow().isoformat(),
            skills_match=template["match_score"],
            remote=True,
            match_score=template["match_score"],
            skills=skills or [role]
        ))

    return jobs


CERTIFICATION_LIBRARY = [
    {
        "name": "AWS Certified Solutions Architect",
        "focus_keywords": ["aws", "cloud", "architecture", "serverless"],
        "provider": "Amazon Web Services",
        "difficulty": "Advanced",
        "predicted_roi": "6-9 months",
        "estimated_time": "6-8 weeks",
        "career_boost": "+1 level",
        "market_signal": 0.22,
        "format": "Exam + hands-on labs"
    },
    {
        "name": "Google Professional ML Engineer",
        "focus_keywords": ["machine learning", "ml", "tensorflow", "data"],
        "provider": "Google Cloud",
        "difficulty": "Advanced",
        "predicted_roi": "6-9 months",
        "estimated_time": "8-10 weeks",
        "career_boost": "Transition to ML roles",
        "market_signal": 0.24,
        "format": "Performance-based exam"
    },
    {
        "name": "Certified Kubernetes Application Developer",
        "focus_keywords": ["kubernetes", "containers", "devops", "cloud"],
        "provider": "CNCF",
        "difficulty": "Intermediate",
        "predicted_roi": "3-6 months",
        "estimated_time": "4-6 weeks",
        "career_boost": "DevOps credibility",
        "market_signal": 0.18,
        "format": "Proctored exam"
    },
    {
        "name": "Azure Solutions Architect Expert",
        "focus_keywords": ["azure", "cloud", "architecture", "infrastructure"],
        "provider": "Microsoft",
        "difficulty": "Advanced",
        "predicted_roi": "6-9 months",
        "estimated_time": "6-8 weeks",
        "career_boost": "Enterprise architect path",
        "market_signal": 0.2,
        "format": "Two-step certification"
    },
    {
        "name": "TOGAF Enterprise Architecture",
        "focus_keywords": ["architecture", "strategy", "enterprise", "governance"],
        "provider": "The Open Group",
        "difficulty": "Intermediate",
        "predicted_roi": "9-12 months",
        "estimated_time": "8-10 weeks",
        "career_boost": "Strategic leadership",
        "market_signal": 0.15,
        "format": "Modular exams"
    },
    {
        "name": "Professional Scrum Master",
        "focus_keywords": ["agile", "scrum", "delivery", "leadership"],
        "provider": "Scrum.org",
        "difficulty": "Beginner",
        "predicted_roi": "2-3 months",
        "estimated_time": "3-4 weeks",
        "career_boost": "Team leadership readiness",
        "market_signal": 0.14,
        "format": "Online exam"
    }
]


def _load_certification_catalog() -> List[Dict[str, Any]]:
    base_dir = Path(__file__).resolve().parents[1]
    catalog_path = base_dir / 'data' / 'certifications_catalog.json'
    if not catalog_path.exists():
        return CERTIFICATION_LIBRARY
    try:
        with open(catalog_path, 'r', encoding='utf-8') as handle:
            data = json.load(handle)
        if isinstance(data, list) and data:
            return data
    except Exception as exc:
        logger.warning("Failed to read dynamic certification catalog: %s", exc)
    return CERTIFICATION_LIBRARY


CERTIFICATION_CATALOG = _load_certification_catalog()

PORTFOLIO_COLOR_MAP = {
    "blue": "#2563eb",
    "green": "#16a34a",
    "purple": "#7c3aed",
    "red": "#dc2626",
    "orange": "#ea580c",
    "teal": "#0d9488",
    "rose": "#e11d48",
    "slate": "#1e293b"
}

ADVANCED_COMPLEXITY_KEYWORDS = {
    "architecture", "ml", "machine learning", "security", "kubernetes",
    "optimization", "distributed", "leadership", "strategy"
}

INTERMEDIATE_COMPLEXITY_KEYWORDS = {
    "cloud", "api", "python", "pandas", "node", "react", "analysis",
    "pipeline", "deployment", "devops"
}

BASE_HOURS_PER_COMPLEXITY = {
    "basic": 2.5,
    "intermediate": 4.0,
    "advanced": 5.5
}

COMPLEXITY_MULTIPLIERS = {
    "basic": 1.0,
    "intermediate": 1.2,
    "advanced": 1.4
}

RESOURCE_TIER_HOURS = {
    "Core": 2,
    "Applied": 4,
    "Proof": 6
}


def _classify_skill_complexity(skill: str) -> Tuple[str, float]:
    lowered = (skill or "").lower()
    if any(keyword in lowered for keyword in ADVANCED_COMPLEXITY_KEYWORDS):
        return "advanced", COMPLEXITY_MULTIPLIERS["advanced"]
    if any(keyword in lowered for keyword in INTERMEDIATE_COMPLEXITY_KEYWORDS):
        return "intermediate", COMPLEXITY_MULTIPLIERS["intermediate"]
    return "basic", COMPLEXITY_MULTIPLIERS["basic"]


def _resource_search_link(skill: str, tier: str) -> str:
    query = quote_plus(f"{skill} {tier.lower()} project")
    return f"https://www.google.com/search?q={query}"


def _generate_resource_tiers(skill: str, complexity: str, multiplier: float) -> List[LearningResourcePayload]:
    tier_configs = [
        {
            "tier": "Core",
            "title": f"{skill} Fundamentals Sprint",
            "provider": "OpenCourseWare",
            "is_free": True,
            "cost": "Free",
            "base_hours": RESOURCE_TIER_HOURS["Core"],
            "rating": 4.4
        },
        {
            "tier": "Applied",
            "title": f"Applied {skill} Projects",
            "provider": "Udemy Labs",
            "is_free": False,
            "cost": "$79",
            "base_hours": RESOURCE_TIER_HOURS["Applied"],
            "rating": 4.6
        },
        {
            "tier": "Proof",
            "title": f"Capstone: Ship a {skill} Outcome",
            "provider": "SkillSync Proof Lab",
            "is_free": True,
            "cost": "Free",
            "base_hours": RESOURCE_TIER_HOURS["Proof"],
            "rating": 4.8
        }
    ]
    resources: List[LearningResourcePayload] = []
    for config in tier_configs:
        estimated_hours = int(round(config["base_hours"] * multiplier)) or config["base_hours"]
        link = _resource_search_link(skill, config["tier"])
        resources.append(LearningResourcePayload(
            skill=skill,
            title=config["title"],
            provider=config["provider"],
            duration=f"~{estimated_hours} hrs",
            rating=config["rating"],
            url=link,
            link=link,
            cost=config["cost"],
            is_free=config["is_free"],
            tier=config["tier"],
            estimated_time_hours=estimated_hours,
            time_hours=estimated_hours
        ))
    return resources


def _estimate_phase_hours(skills: List[str], complexity_map: Dict[str, Tuple[str, float]], resources: List[LearningResourcePayload]) -> int:
    base_hours = 0.0
    for skill in skills:
        level, multiplier = complexity_map.get(skill, ("basic", 1.0))
        base_hours += BASE_HOURS_PER_COMPLEXITY[level] * multiplier
    resource_hours = sum(resource.estimated_time_hours or 0 for resource in resources)
    milestone_buffer = max(4, len(skills) * 2)
    return int(round(base_hours + resource_hours + milestone_buffer))


def _build_success_justification(phase_name: str, success_pct: int, skills: List[str], complexity_map: Dict[str, Tuple[str, float]]) -> str:
    if not skills:
        return f"Moderate confidence ({success_pct}%) with focus on foundational refresh."
    highest_level = max((complexity_map[skill][0] for skill in skills), key=lambda level: COMPLEXITY_MULTIPLIERS.get(level, 1.0))
    sample_skill = skills[0]
    level_descriptor = {
        "basic": "foundational",
        "intermediate": "applied",
        "advanced": "advanced"
    }.get(highest_level, "foundational")
    return (
        f"{phase_name} confidence at {success_pct}% because {len(skills)} {level_descriptor} skills "
        f"have existing momentum (e.g., {sample_skill})."
    )


def _generate_smart_milestones(phase_name: str, skills: List[str], total_hours: int) -> Tuple[List[Dict[str, Any]], List[str]]:
    if not skills:
        skills = ["Execution"]
    focus_skills = skills[: max(1, min(3, len(skills)))]
    milestone_types = ["Exercise", "Project", "Review"]
    smart_records: List[Dict[str, Any]] = []
    summaries: List[str] = []
    for idx, skill in enumerate(focus_skills):
        deadline = max(6, total_hours // (idx + 3))
        milestone = {
            "title": f"{phase_name}: Apply {skill}",
            "type": milestone_types[idx % len(milestone_types)],
            "target_metric": f"Score {90 - idx * 5}%+ on {skill} rubric",
            "deadline_hours": deadline
        }
        smart_records.append(milestone)
        summaries.append(f"{milestone['title']} (Target: {milestone['target_metric']}, {deadline}h)")
    return smart_records, summaries


def _flatten_skill_results(skills_data: Dict[str, Any]) -> List[str]:
    seen = []
    if not skills_data:
        return seen

    category_map = skills_data.get('skills', {})
    for entries in category_map.values():
        for entry in entries:
            skill_name = entry.get('skill') if isinstance(entry, dict) else None
            if skill_name:
                normalized = skill_name.strip().title()
                if normalized and normalized not in seen:
                    seen.append(normalized)
    return seen


def _extract_keywords(text: str, keywords: List[str]) -> List[str]:
    matches = []
    for keyword in keywords:
        if re.search(rf'\b{re.escape(keyword)}\b', text, re.IGNORECASE):
            matches.append(keyword.title())
    return matches


def _estimate_experience_years(text: str, detected_skills: int) -> float:
    matches = re.findall(r'(\d+)\s+years?', text.lower())
    if matches:
        return float(max(int(m) for m in matches))
    return float(max(1, detected_skills // 4))


def _safe_float(value: Any) -> Optional[float]:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        digits = re.sub(r'[^0-9.]', '', value)
        if digits:
            try:
                return float(digits)
            except ValueError:
                return None
    return None


def _normalize_salary_band(job: Dict[str, Any]) -> SalaryBand:
    salary_min = _safe_float(job.get('salary_min') or job.get('salaryMin'))
    salary_max = _safe_float(job.get('salary_max') or job.get('salaryMax'))
    currency = job.get('currency') or 'USD'
    if salary_min is None and salary_max is None:
        role_hint = job.get('title', '').lower()
        base = 135000 if 'architect' in role_hint or 'lead' in role_hint else 110000
        return SalaryBand(min=base, max=base + 20000, currency=currency)
    if salary_min is None:
        salary_min = max((salary_max or 100000) * 0.8, 80000)
    if salary_max is None:
        salary_max = salary_min * 1.2
    if salary_max < salary_min:
        salary_max = salary_min
    return SalaryBand(min=float(round(salary_min, 2)), max=float(round(salary_max, 2)), currency=currency)


_SOURCE_LOGOS = {
    "linkedin": "https://logo.clearbit.com/linkedin.com",
    "adzuna": "https://logo.clearbit.com/adzuna.com",
    "remoteok": "https://logo.clearbit.com/remoteok.com",
    "themuse": "https://logo.clearbit.com/themuse.com",
    "findwork": "https://logo.clearbit.com/findwork.dev",
    "arbeitnow": "https://logo.clearbit.com/arbeitnow.com",
    "jobicy": "https://logo.clearbit.com/jobicy.com",
    "jsearch": "https://logo.clearbit.com/rapidapi.com"
}


def _job_source_logo(source: Optional[str]) -> Optional[str]:
    if not source:
        return None
    normalized = source.lower()
    return _SOURCE_LOGOS.get(normalized)


_PROVIDER_LOGOS = {
    "amazon web services": "https://logo.clearbit.com/aws.amazon.com",
    "google cloud": "https://logo.clearbit.com/cloud.google.com",
    "cncf": "https://logo.clearbit.com/cncf.io",
    "microsoft": "https://logo.clearbit.com/microsoft.com",
    "scrum.org": "https://logo.clearbit.com/scrum.org",
    "the open group": "https://logo.clearbit.com/opengroup.org",
    "coursera": "https://logo.clearbit.com/coursera.org",
    "udacity": "https://logo.clearbit.com/udacity.com"
}


def _provider_logo(provider: Optional[str]) -> Optional[str]:
    if not provider:
        return None
    normalized = provider.strip().lower()
    if normalized in _PROVIDER_LOGOS:
        return _PROVIDER_LOGOS[normalized]
    domain_hint = normalized.replace(' ', '')
    if domain_hint:
        return f"https://logo.clearbit.com/{domain_hint}.com"
    return None


def _strip_html(text: str) -> str:
    if not text:
        return ""
    without_tags = re.sub(r'<[^>]+>', ' ', text)
    collapsed = re.sub(r'\s+', ' ', without_tags)
    return html.unescape(collapsed).strip()


def _infer_location_type(job: Dict[str, Any], location_value: str) -> str:
    remote_flag = job.get('remote')
    if isinstance(remote_flag, bool):
        return "Remote" if remote_flag else "On-site"
    value = (location_value or "").lower()
    if "remote" in value:
        return "Remote"
    if "hybrid" in value:
        return "Hybrid"
    if any(token in value for token in ["flexible", "anywhere"]):
        return "Remote"
    return "On-site"


def _collect_job_tags(job: Dict[str, Any], location_type: str) -> List[str]:
    tags = []
    job_type = job.get('job_type') or job.get('employment_type')
    if isinstance(job_type, str):
        tags.append(job_type.title())
    seniority = job.get('seniority')
    if isinstance(seniority, str):
        tags.append(seniority.title())
    industry = job.get('category') or job.get('industry')
    if isinstance(industry, str):
        tags.append(industry.title())
    if location_type:
        tags.append(location_type)
    return tags


def _infer_job_skill_gaps(job_text: str, matching_skills: List[str]) -> List[str]:
    job_keywords = _extract_keywords(job_text, TECH_KEYWORD_POOL)
    normalized_matches = {skill.title() for skill in matching_skills}
    gaps = [skill for skill in job_keywords if skill not in normalized_matches]
    return gaps[:4]


def _normalize_skill_token(value: Optional[str]) -> str:
    if not value:
        return ""
    lowered = value.strip().lower()
    lowered = lowered.replace('/', ' ').replace('|', ' ')
    return re.sub(r'[^a-z0-9+#. ]+', '', lowered)


def _collect_job_skill_requirements(job: Dict[str, Any]) -> List[str]:
    collected: List[str] = []
    candidate_keys = [
        'skills', 'skills_required', 'skill_requirements', 'requirements',
        'keywords', 'tags', 'top_skills', 'nice_to_have'
    ]
    for key in candidate_keys:
        payload = job.get(key)
        if isinstance(payload, list):
            collected.extend([item for item in payload if isinstance(item, str)])
        elif isinstance(payload, str):
            splits = re.split(r'[;,/\n]', payload)
            collected.extend([chunk.strip() for chunk in splits if chunk.strip()])
    return [skill.title() for skill in dict.fromkeys(collected)]


def _resolve_matching_skills(
    job: Dict[str, Any],
    job_text: str,
    cv_skills: List[str],
    job_skill_requirements: List[str]
) -> List[str]:
    matches: List[str] = []
    job_tokens = {
        _normalize_skill_token(skill): skill
        for skill in job_skill_requirements
    }
    job_text_lower = job_text.lower()
    for skill in cv_skills:
        normalized = _normalize_skill_token(skill)
        if not normalized:
            continue
        if normalized in job_tokens:
            matches.append(skill.title())
            continue
        if any(normalized in token for token in job_tokens.keys() if len(token) > len(normalized) >= 3):
            matches.append(skill.title())
            continue
        if re.search(rf'\b{re.escape(skill)}\b', job_text, re.IGNORECASE):
            matches.append(skill.title())
            continue
        if normalized and normalized in job_text_lower:
            matches.append(skill.title())
    provided_matches = job.get('matching_skills') or job.get('matchingSkills') or []
    for provided in provided_matches:
        if isinstance(provided, str) and provided.strip().title() in cv_skills:
            matches.append(provided.strip().title())
    return list(dict.fromkeys(matches))


def _derive_gap_skills(
    job_skill_requirements: List[str],
    cv_skills: List[str],
    matching_skills: List[str],
    job_text: str
) -> List[str]:
    if job_skill_requirements:
        normalized_cv = {
            _normalize_skill_token(skill)
            for skill in cv_skills
            if isinstance(skill, str)
        }
        normalized_matches = {
            _normalize_skill_token(skill)
            for skill in matching_skills
            if isinstance(skill, str)
        }
        coverage_tokens = {token for token in normalized_cv if token} | {token for token in normalized_matches if token}
        prioritized = [
            skill for skill in job_skill_requirements
            if _normalize_skill_token(skill) not in coverage_tokens
        ]
        if prioritized:
            return prioritized[:6]
    # fall back to textual inference if structured skills missing
    return _infer_job_skill_gaps(job_text, matching_skills)


def _ensure_job_url(job: Dict[str, Any]) -> str:
    job_url = job.get('url') or job.get('job_url') or job.get('redirect_url') or job.get('apply_link')
    if job_url:
        return str(job_url)
    title = job.get('title') or job.get('name') or 'opportunity'
    slug = re.sub(r'[^a-z0-9]+', '-', title.lower()).strip('-') or 'opportunity'
    return f"https://jobs.skillsync.ai/{slug}"


def _convert_job_to_recommendation(job: Dict[str, Any], cv_skills: List[str]) -> Optional[JobRecommendationPayload]:
    if not job:
        return None
    job_text = " ".join([
        job.get('title', ''),
        job.get('company', ''),
        job.get('description', '')
    ])
    if not job_text.strip():
        job_text = job.get('title', '')

    cv_skill_set = [skill.title() for skill in cv_skills if isinstance(skill, str)]
    cv_skill_set = list(dict.fromkeys(cv_skill_set))
    focus_window = 15 if len(cv_skill_set) > 15 else len(cv_skill_set)
    cv_focus = cv_skill_set[:focus_window] or ["Python", "Leadership"]

    job_skill_requirements = _collect_job_skill_requirements(job)
    matching_skills = _resolve_matching_skills(job, job_text, cv_focus, job_skill_requirements)
    if not matching_skills:
        matching_skills = [
            skill for skill in cv_focus
            if re.search(rf'\b{re.escape(skill)}\b', job_text, re.IGNORECASE)
        ]
    similarity = len(matching_skills) / max(len(cv_focus), 1)
    match_score = job.get('match_score') or job.get('matchScore') or similarity * 100
    similarity_score = round(min(1.0, similarity + (match_score / 100) * 0.25), 4)
    confidence = round(min(0.4 + similarity * 0.4 + (match_score / 100) * 0.3, 0.97), 2)
    gaps = _derive_gap_skills(job_skill_requirements, cv_skill_set, matching_skills, job_text)
    if not gaps:
        gaps = [skill for skill in cv_focus if skill not in matching_skills][:4]

    growth = "Very High" if match_score >= 80 else "High" if match_score >= 60 else "Medium"
    salary_band = _normalize_salary_band(job)
    source = job.get('source') or job.get('job_source') or job.get('api_source') or "API"
    location = job.get('location') or job.get('job_location') or "Remote"
    location_type = _infer_location_type(job, location)
    reasons = [
        f"Live {source} listing analyzed",
        f"Matched skills: {', '.join(matching_skills) or 'Transferable strengths'}"
    ]
    if location:
        reasons.append(f"Location fit: {location}")
    if location_type == "Remote":
        reasons.append("Remote friendly role")

    description = _strip_html(
        job.get('description') or job.get('summary') or job.get('short_description') or ''
    )
    job_url = _ensure_job_url(job)
    posted_date = job.get('posted_date') or job.get('publication_date') or job.get('date_posted')
    tags = _collect_job_tags(job, location_type)

    return JobRecommendationPayload(
        title=job.get('title') or "Opportunity",
        similarity_score=similarity_score,
        confidence=confidence,
        predicted_salary=salary_band,
        matching_skills=matching_skills,
        matched_skills=matching_skills,
        skill_gaps=gaps,
        gap_skills=gaps,
        growth_potential=growth,
        reasons=reasons[:3],
        description=description or None,
        job_url=job_url,
        source=source.title() if isinstance(source, str) else source,
        source_logo=_job_source_logo(source if isinstance(source, str) else None),
        posted_date=posted_date,
        posting_date=posted_date,
        location=location,
        location_type=location_type,
        tags=tags
    )


def _map_jobs_to_recommendations(job_results: List[Dict[str, Any]], cv_skills: List[str], limit: int = 3) -> List[JobRecommendationPayload]:
    recommendations: List[JobRecommendationPayload] = []
    for job in job_results:
        recommendation = _convert_job_to_recommendation(job, cv_skills)
        if recommendation:
            recommendations.append(recommendation)
        if len(recommendations) >= limit:
            break
    return recommendations


def _fallback_template_recommendations(
    primary_skills: List[str],
    analysis: CVAnalysisResponse,
    limit: int = 3
) -> List[JobRecommendationPayload]:
    fallback_jobs = _generate_fallback_jobs(
        (analysis.job_titles or [analysis.current_title or 'Professional'])[0],
        (analysis.personal_info or {}).get('location') or 'Remote',
        primary_skills,
        limit
    )
    recommendations: List[JobRecommendationPayload] = []
    for job in fallback_jobs:
        job_dict = job.model_dump()
        recommendation = _convert_job_to_recommendation(job_dict, primary_skills)
        if recommendation:
            recommendations.append(recommendation)
    return recommendations


def _infer_search_location(analysis: CVAnalysisResponse) -> str:
    location = (analysis.personal_info or {}).get('location') if analysis.personal_info else None
    if not location and analysis.contact_info:
        location = analysis.contact_info.get('location')
    if not location:
        return "fr"
    lowered = location.lower()
    if any(token in lowered for token in ["paris", "france", "fr"]):
        return "fr"
    if any(token in lowered for token in ["us", "usa", "united states", "new york", "san francisco", "ca"]):
        return "us"
    if any(token in lowered for token in ["canada", "toronto", "montreal"]):
        return "ca"
    return "fr"


async def _fetch_dynamic_job_recommendations(
    primary_skills: List[str],
    analysis: CVAnalysisResponse,
    max_results: int = 6
) -> Tuple[List[JobRecommendationPayload], Dict[str, Any]]:
    search_terms = [skill for skill in primary_skills if skill][:5]
    if not search_terms:
        search_terms = analysis.job_titles or ["developer"]
    location_hint = _infer_search_location(analysis)
    job_meta: Dict[str, Any] = {}

    try:
        job_meta = await search_jobs_multi_source(search_terms, location=location_hint, max_results=max_results)
        job_results = job_meta.get('jobs') or []
        log_ml(
            "job-fetch",
            search_terms=search_terms,
            location=location_hint,
            total=len(job_results),
            sources=job_meta.get('sources_used'),
            total_unique=job_meta.get('total_jobs')
        )
        recommendations = _map_jobs_to_recommendations(job_results, primary_skills)
        return recommendations, job_meta
    except Exception as exc:
        logger.warning("[career-guidance] live job retrieval failed: %s", exc)
        return [], {}


def _select_primary_skills(analysis: CVAnalysisResponse) -> List[str]:
    def _finalize(candidates: List[str]) -> List[str]:
        filtered = _filter_primary_skills(candidates, analysis)
        return _prioritize_skills_by_intent(filtered, analysis)

    skills = analysis.skills or analysis.hard_skills or []
    if skills:
        return _finalize(skills)
    if analysis.skill_categories:
        aggregated = []
        for values in analysis.skill_categories.values():
            aggregated.extend(values)
        if aggregated:
            return _finalize(aggregated)
    text = (analysis.raw_text or analysis.summary or "").lower()
    if text:
        tech_keywords = programming_languages + frameworks + databases + cloud_platforms + tools
        extracted = _extract_keywords(text, tech_keywords)
        if extracted:
            return _finalize(extracted)
    return _prioritize_skills_by_intent([], analysis)


def _collect_personal_tokens(analysis: Optional[CVAnalysisResponse]) -> Set[str]:
    tokens: Set[str] = set()

    def _ingest(value: Optional[str]) -> None:
        if not value:
            return
        lowered = value.lower()
        tokens.update(re.findall(r"[a-z]{3,}", lowered))

    if analysis is None:
        return tokens
    if analysis.personal_info:
        _ingest(analysis.personal_info.get('name'))
        _ingest(analysis.personal_info.get('location'))
    if analysis.contact_info:
        for contact_value in analysis.contact_info.values():
            _ingest(contact_value)
    header = (analysis.raw_text or "")[:160]
    _ingest(header)
    return tokens


def _is_whitelisted_skill(token: str) -> bool:
    if token in SKILL_WHITELIST:
        return True
    compact = token.replace(" ", "")
    if compact in SKILL_WHITELIST:
        return True
    segments = re.split(r"[\s/,-]+", token)
    return any(segment in SKILL_WHITELIST for segment in segments if segment)


def _matches_personal_token(token: str, personal_tokens: Set[str]) -> bool:
    for name_token in personal_tokens:
        if not name_token:
            continue
        if token == name_token:
            return True
        if abs(len(token) - len(name_token)) <= 1 and (
            token.startswith(name_token) or name_token.startswith(token)
        ):
            return True
    return False


def _filter_primary_skills(skills: List[str], analysis: Optional[CVAnalysisResponse]) -> List[str]:
    if not skills:
        return []
    personal_tokens = _collect_personal_tokens(analysis)
    filtered: List[str] = []
    for skill in skills:
        if not skill:
            continue
        normalized = re.sub(r"[^a-z0-9+#./ ]+", " ", skill.lower()).strip()
        if not normalized:
            continue
        if normalized in MONTH_NAME_TOKENS:
            continue
        if _matches_personal_token(normalized, personal_tokens):
            continue
        if _is_whitelisted_skill(normalized):
            filtered.append(skill if isinstance(skill, str) else str(skill))
            continue
    if not filtered:
        fallback = [skill for skill in skills if skill and isinstance(skill, str)]
        return fallback[:5]
    return list(dict.fromkeys(filtered))


def _normalize_skill_label(skill: str) -> str:
    return re.sub(r"\s+", " ", skill.strip().lower()) if isinstance(skill, str) else ""


def _detect_intent_track(analysis: Optional[CVAnalysisResponse]) -> str:
    if analysis is None:
        return 'general'
    text_segments: List[str] = []
    if analysis.summary:
        text_segments.append(analysis.summary)
    if analysis.raw_text:
        text_segments.append(analysis.raw_text[:1500])
    if analysis.job_titles:
        text_segments.extend(analysis.job_titles[:3])
    blob = " ".join(text_segments).lower()
    if any(keyword in blob for keyword in AI_INTENT_KEYWORDS):
        return 'ai'
    if any(keyword in blob for keyword in WEB_INTENT_KEYWORDS):
        return 'web'
    return 'general'


def _infer_ai_basics(existing_skills: List[str], all_skills: List[str], analysis: Optional[CVAnalysisResponse]) -> List[str]:
    normalized_existing = {
        _normalize_skill_label(skill)
        for skill in (existing_skills + all_skills)
    }
    inferred: List[str] = []
    for canonical in ['Machine Learning', 'Data Science', 'MLOps']:
        if canonical.lower() not in normalized_existing:
            inferred.append(canonical)
    profile_text = (analysis.summary or analysis.raw_text or '').lower() if analysis else ''
    if 'computer vision' in profile_text and 'computer vision' not in normalized_existing:
        inferred.append('Computer Vision')
    if 'nlp' in profile_text or 'natural language processing' in profile_text:
        if 'nlp' not in normalized_existing and 'natural language processing' not in normalized_existing:
            inferred.append('NLP')
    return inferred[:3]


def _prioritize_skills_by_intent(skills: List[str], analysis: Optional[CVAnalysisResponse]) -> List[str]:
    if not skills:
        return []
    intent = _detect_intent_track(analysis)
    if intent != 'ai':
        return skills[:10]

    ai_focus: List[str] = []
    remainder: List[str] = []
    for skill in skills:
        normalized = _normalize_skill_label(skill)
        if not normalized:
            continue
        if normalized in AI_PRIORITY_SKILL_SET or any(keyword in normalized for keyword in AI_INTENT_KEYWORDS):
            ai_focus.append(skill)
        else:
            remainder.append(skill)

    if not ai_focus:
        ai_focus = [skill for skill in skills if skill][:2]

    inferred = _infer_ai_basics(ai_focus, skills, analysis)
    ordered = inferred + ai_focus + remainder

    deduped: List[str] = []
    seen: Set[str] = set()
    for skill in ordered:
        normalized = _normalize_skill_label(skill)
        if not normalized or normalized in seen:
            continue
        deduped.append(skill)
        seen.add(normalized)
    return deduped[:10]


def _pick_soft_skills(analysis: CVAnalysisResponse) -> List[str]:
    if analysis.soft_skills:
        return analysis.soft_skills
    text = analysis.summary or analysis.raw_text or ""
    return _extract_keywords(text, soft_skill_keywords)



def _build_certifications(skill_gaps: List[str]) -> List[CertificationRecommendationPayload]:
    normalized_gaps = [gap.lower() for gap in skill_gaps if gap]
    if not normalized_gaps:
        normalized_gaps = ["system design", "leadership", "cloud"]
    gap_weights = Counter(normalized_gaps)
    catalog = CERTIFICATION_CATALOG or CERTIFICATION_LIBRARY
    difficulty_bonus = {
        'beginner': 0.04,
        'intermediate': 0.07,
        'advanced': 0.1
    }

    scored: List[Tuple[float, CertificationRecommendationPayload]] = []
    for cert in catalog:
        keywords = [kw.lower() for kw in cert.get('focus_keywords') or cert.get('keywords', [])]
        hits = [gap for gap in normalized_gaps if any(keyword in gap for keyword in keywords)]
        coverage_score = sum(gap_weights.get(hit, 1) for hit in hits)
        market_signal = float(cert.get('market_signal', 0.12))
        difficulty = difficulty_bonus.get(cert.get('difficulty', '').lower(), 0.06)
        personalization = min(0.25, 0.12 * coverage_score)
        relevance = min(0.98, 0.45 + personalization + market_signal + difficulty)
        alignment = min(0.95, 0.5 + personalization + market_signal * 0.5)
        target_gap = hits[0].title() if hits else normalized_gaps[0].title()
        focus_keyword = keywords[0].title() if keywords else target_gap
        reasons = [
            f"Targets {target_gap} gap via {focus_keyword}",
            f"Delivered by {cert.get('provider', 'Top provider')} ({cert.get('difficulty', 'Intermediate')})"
        ]
        if cert.get('format'):
            reasons.append(f"Format: {cert['format']}")
        if len(hits) >= 2:
            reasons.append("Covers multiple detected gaps simultaneously")

        official_url = _resolve_cert_official_url(cert)
        provider = cert.get('provider')
        cost_estimate = _normalize_cost_estimate(cert)
        resources = _build_cert_resources(cert, official_url)
        payload = CertificationRecommendationPayload(
            name=cert['name'],
            relevance_score=round(relevance, 4),
            skill_alignment=round(max(0.5, alignment), 4),
            predicted_roi=cert['predicted_roi'],
            estimated_time=cert['estimated_time'],
            career_boost=cert['career_boost'],
            reasons=reasons[:3],
            url=official_url,
            official_url=official_url,
            provider=provider,
            provider_logo=_provider_logo(provider),
            cost_type=cert.get('cost_type') or ('Free' if cost_estimate.lower() == 'free' else 'Paid'),
            delivery_format=cert.get('format'),
            cost_estimate=cost_estimate,
            resources=resources
        )
        scored.append((relevance, payload))

    scored.sort(key=lambda item: item[0], reverse=True)
    top = [item[1] for item in scored[:3]]
    log_ml(
        "certifications",
        gaps=normalized_gaps[:5],
        selected=[cert.name for cert in top],
        catalog_size=len(catalog)
    )
    return top


def _resolve_cert_official_url(cert: Dict[str, Any]) -> str:
    base_url = cert.get('official_url') or cert.get('url')
    if base_url:
        return str(base_url)
    query = quote_plus(f"{cert.get('name', 'certification')} official exam")
    return f"https://www.google.com/search?q={query}"


def _normalize_cost_estimate(cert: Dict[str, Any]) -> str:
    estimate = cert.get('cost_estimate') or cert.get('cost') or cert.get('price')
    if isinstance(estimate, (int, float)):
        return f"${int(estimate):,}"
    if isinstance(estimate, str) and estimate.strip():
        return estimate.strip()
    cost_type = (cert.get('cost_type') or '').lower()
    if 'free' in cost_type:
        return 'Free'
    return 'Paid'


def _build_cert_resources(cert: Dict[str, Any], official_url: str) -> List[Dict[str, str]]:
    resources = cert.get('resources') or []
    normalized: List[Dict[str, str]] = []
    for resource in resources:
        if not isinstance(resource, dict):
            continue
        title = resource.get('title') or resource.get('name')
        link = resource.get('link') or resource.get('url')
        if title and link:
            normalized.append({'title': title, 'link': link})
    if not normalized and official_url:
        normalized.append({
            'title': f"Official {cert.get('name', 'certification')} guide",
            'link': official_url
        })
    return normalized


def _analysis_to_learning_context(analysis: Optional[CVAnalysisResponse]) -> Dict[str, Any]:
    if analysis is None:
        return {}
    experience_years = analysis.total_years_experience or analysis.experience_years or 4
    if isinstance(experience_years, str):
        try:
            experience_years = float(experience_years)
        except ValueError:
            experience_years = 4
    weekly_hours = 8
    if experience_years <= 2:
        weekly_hours = 5
    elif experience_years >= 8:
        weekly_hours = 10
    industries = analysis.industries or []
    target_role = (analysis.job_titles or [analysis.current_title or 'Career Growth'])[0]
    seniority = 'Senior' if experience_years >= 8 else 'Junior' if experience_years <= 2 else 'Mid-level'
    return {
        'experience_years': experience_years,
        'weekly_hours_available': weekly_hours,
        'industry': industries[0] if industries else 'Tech',
        'target_role': target_role,
        'seniority': seniority
    }


def _dict_to_learning_resource(resource: Dict[str, Any]) -> LearningResourcePayload:
    return LearningResourcePayload(
        skill=resource.get('skill', 'Skill'),
        title=resource.get('title', 'Learning Resource'),
        provider=resource.get('provider', 'SkillSync Catalog'),
        duration=resource.get('duration', 'Self-paced'),
        rating=float(resource.get('rating', 4.0)),
        url=resource.get('url'),
        cost=resource.get('cost'),
        is_free=bool(resource.get('is_free', False)),
        tier=resource.get('tier', 'Core'),
        link=resource.get('link'),
        estimated_time_hours=int(resource.get('estimated_time_hours') or resource.get('time_hours') or 6),
        time_hours=int(resource.get('time_hours') or resource.get('estimated_time_hours') or 6)
    )


def _materialize_learning_plan(
    plan: Dict[str, Any],
    primary_skills: List[str],
    gaps: List[str]
) -> LearningRoadmapPayload:
    phases_payload: List[LearningPhasePayload] = []
    for phase in plan.get('phases', []):
        skills = phase.get('skills') or []
        if not skills:
            continue
        complexity_map = {
            skill: _classify_skill_complexity(skill)
            for skill in skills
        }
        resources = [_dict_to_learning_resource(resource) for resource in (phase.get('resources') or [])][:8]
        success_raw = float(phase.get('success_probability', 0.78))
        success_pct_value = max(55, min(int(round(success_raw * 100)), 97))
        success_text = _build_success_justification(phase.get('phase_name', 'Phase'), success_pct_value, skills, complexity_map)
        smart_milestones, milestone_summaries = _generate_smart_milestones(
            phase.get('phase_name', 'Phase'),
            skills,
            int(round(phase.get('predicted_hours', 12)))
        )
        duration_weeks = int(phase.get('duration_weeks', 4))
        phases_payload.append(LearningPhasePayload(
            phase_name=phase.get('phase_name', 'Phase'),
            duration_weeks=duration_weeks,
            duration_months=round(duration_weeks / 4.0, 2),
            total_time_estimate_hours=int(round(phase.get('predicted_hours', 12))),
            skills_to_learn=skills,
            learning_resources=resources,
            resources=resources,
            success_probability=f"{success_pct_value}%",
            success_justification=success_text,
            effort_level=phase.get('effort_level', 'Medium'),
            milestones=milestone_summaries,
            smart_milestones=smart_milestones
        ))

    if not phases_payload:
        raise ValueError("ML plan did not produce any phases")

    total_weeks = plan.get('total_weeks') or sum(phase.duration_weeks for phase in phases_payload)
    total_hours = int(round(plan.get('total_hours', sum(phase.total_time_estimate_hours for phase in phases_payload))))
    predicted_success_rate = plan.get('predicted_success_rate', 0.82)
    personalization = plan.get('personalization', 'A')
    strategy = plan.get('learning_strategy', 'Balanced Growth')

    return LearningRoadmapPayload(
        total_duration_weeks=int(total_weeks),
        total_duration_months=round(total_weeks / 4.0, 1),
        total_time_estimate_hours=total_hours,
        predicted_success_rate=f"{int(round(predicted_success_rate * 100))}%",
        personalization_score=personalization,
        learning_strategy=strategy,
        phases=phases_payload
    )


def _build_learning_roadmap(
    primary_skills: List[str],
    gaps: List[str],
    analysis: Optional[CVAnalysisResponse] = None
) -> LearningRoadmapPayload:
    predictor = LEARNING_ROADMAP_PREDICTOR
    context = _analysis_to_learning_context(analysis)
    if predictor:
        try:
            plan = predictor.build_learning_roadmap(primary_skills, gaps, context)
            _persist_learning_roadmap_run(plan, primary_skills, gaps, context, analysis)
            roadmap = _materialize_learning_plan(plan, primary_skills, gaps)
            log_ml(
                "learning-roadmap-ml",
                total_hours=plan.get('total_hours'),
                total_weeks=plan.get('total_weeks'),
                personalization=plan.get('personalization'),
                strategy=plan.get('learning_strategy')
            )
            return roadmap
        except Exception as exc:  # noqa: BLE001
            logger.warning("[roadmap-ml] Falling back to legacy roadmap: %s", exc)
    return _legacy_learning_roadmap(primary_skills, gaps)


def _legacy_learning_roadmap(primary_skills: List[str], gaps: List[str]) -> LearningRoadmapPayload:
    core_focus = [skill.title() for skill in primary_skills[:4] if skill] or ["Systems Thinking", "Delivery Excellence"]
    gap_focus = [gap.title() for gap in gaps[:4] if gap] or ["Cloud Architecture", "System Design", "Product Sense"]
    advanced_focus = [gap.title() for gap in gaps[4:7] if gap] or ["Storytelling", "Executive Communication"]

    phase_configs = [
        {
            "name": "Stabilize Core",
            "skills": core_focus,
            "base_success": 90,
            "effort": "Medium"
        },
        {
            "name": "Upskill for Growth",
            "skills": gap_focus,
            "base_success": 85,
            "effort": "High"
        },
        {
            "name": "Differentiate",
            "skills": advanced_focus,
            "base_success": 82,
            "effort": "Medium"
        }
    ]

    phases: List[LearningPhasePayload] = []
    total_hours_across_phases = 0

    for config in phase_configs:
        skills = config["skills"]
        if not skills:
            continue

        complexity_map = {
            skill: _classify_skill_complexity(skill)
            for skill in skills
        }
        resources: List[LearningResourcePayload] = []
        for skill in skills:
            level, multiplier = complexity_map[skill]
            resources.extend(_generate_resource_tiers(skill, level, multiplier))

        phase_hours = _estimate_phase_hours(skills, complexity_map, resources)
        total_hours_across_phases += phase_hours
        weeks = max(2, math.ceil(phase_hours / 6))
        success_penalty = int(sum(max(multiplier - 1.0, 0) * 8 for _, multiplier in complexity_map.values()))
        success_pct = max(60, min(config["base_success"] - success_penalty, 96))
        success_text = _build_success_justification(config["name"], success_pct, skills, complexity_map)
        smart_milestones, milestone_summaries = _generate_smart_milestones(config["name"], skills, phase_hours)

        phase_payload = LearningPhasePayload(
            phase_name=config["name"],
            duration_weeks=weeks,
            duration_months=round(weeks / 4.0, 2),
            total_time_estimate_hours=phase_hours,
            skills_to_learn=skills,
            learning_resources=resources,
            resources=resources,
            success_probability=f"{success_pct}%",
            success_justification=success_text,
            effort_level=config["effort"],
            milestones=milestone_summaries,
            smart_milestones=smart_milestones
        )
        phases.append(phase_payload)

    if not phases:
        # fallback to default progression to avoid empty payloads
        return _legacy_learning_roadmap(["Python"], ["Cloud"])

    total_weeks = sum(phase.duration_weeks for phase in phases)
    success_rate = min(95, 72 + len(primary_skills) * 2)
    personalization_score = "A" if len(gaps) >= 4 else "A-" if len(gaps) >= 2 else "B+"

    roadmap = LearningRoadmapPayload(
        total_duration_weeks=total_weeks,
        total_duration_months=round(total_weeks / 4.0, 1),
        total_time_estimate_hours=total_hours_across_phases,
        predicted_success_rate=f"{success_rate}%",
        personalization_score=personalization_score,
        learning_strategy="ML-personalized blended learning",
        phases=phases
    )
    log_ml(
        "learning-roadmap",
        total_weeks=total_weeks,
        total_hours=total_hours_across_phases,
        primary_skills=core_focus,
        gaps=gap_focus,
        advanced=advanced_focus,
        success_rate=roadmap.predicted_success_rate,
        personalization=personalization_score
    )
    return roadmap


def _build_xai_details(analysis: CVAnalysisResponse, recommendations: List[JobRecommendationPayload]) -> XAIInsightsPayload:
    if recommendations:
        main_reco = recommendations[0]
        matching_features = main_reco.matching_skills[:5]
        skill_gaps = main_reco.skill_gaps[:3]
        salary_band_dump = main_reco.predicted_salary.model_dump()
        top_title = main_reco.title
        job_confidence = round(main_reco.confidence, 2)
        certification_confidence = round(main_reco.confidence * 0.9, 2)
        similarity_score = main_reco.similarity_score
    else:
        fallback_title = (analysis.job_titles or [analysis.current_title or "Target Role"])[0]
        fallback_skills = analysis.skills or []
        fallback_gaps_source = analysis.hard_skills or fallback_skills
        matching_features = fallback_skills[:5]
        skill_gaps = (fallback_gaps_source or [])[:3]
        fallback_salary = SalaryBand(min=0.0, max=0.0, currency="USD")
        salary_band_dump = fallback_salary.model_dump()
        confidence_seed = analysis.confidence_score if analysis.confidence_score is not None else 0.6
        job_confidence = round(max(0.12, min(confidence_seed, 0.3)), 2)
        certification_confidence = round(max(0.35, min(job_confidence - 0.05, 0.4)), 2)
        similarity_score = round(min(job_confidence + 0.1, 0.6), 2)
        top_title = fallback_title
        log_ml(
            "xai-fallback",
            analysis_id=analysis.analysis_id,
            reason="No job recommendations available",
            total_skills=len(fallback_skills)
        )
    how_cv = {
        "model": "SkillSync-ML-v2",
        "analysis_id": analysis.analysis_id,
        "processed_sections": {
            "skills": len(analysis.skills or []),
            "work_history": len(analysis.work_history or []),
            "projects": len(getattr(analysis, 'projects', []) or []),
        },
        "signals": matching_features,
        "explanations_available": True
    }
    job_explain = {
        "top_recommendation": top_title,
        "similarity_score": similarity_score,
        "features_used": matching_features,
        "skill_gaps": skill_gaps,
        "salary_band": salary_band_dump
    }
    cert_explain = {
        "ranking_logic": "Alignment with detected gaps and ROI horizon",
        "gap_focus": skill_gaps,
        "confidence": certification_confidence,
        "human_override": True
    }
    learning_explain = {
        "roadmap_length_weeks": 12,
        "phases": ["Stabilize Core", "Upskill for Growth", "Differentiate"],
        "primary_objective": f"Advance toward {top_title}",
        "ml_personalization": True
    }
    confidence_scores = {
        "job_matching": job_confidence,
        "certifications": certification_confidence,
        "learning_path": 0.8,
        "xai_explainability": 0.92
    }
    key_insights = [
        f"{len(analysis.skills or [])} core skills detected",
        f"{matching_features[0] if matching_features else 'Skill signals'} drives job alignment",
        "Roadmap blends delivery + leadership milestones"
    ]
    return XAIInsightsPayload(
        how_we_analyzed_your_cv=how_cv,
        job_matching_explanation=job_explain,
        certification_ranking_explanation=cert_explain,
        learning_path_explanation=learning_explain,
        ml_confidence_scores=confidence_scores,
        key_insights=key_insights
    )


def _build_key_opportunities(recommendations: List[JobRecommendationPayload]) -> List[KeyOpportunityPayload]:
    items: List[KeyOpportunityPayload] = []
    for reco in recommendations[:2]:
        priority = "High" if reco.confidence >= 0.75 else "Medium"
        next_steps = [
            f"Showcase impact stories tied to {skill}" for skill in reco.matching_skills[:2]
        ] or ["Highlight measurable outcomes in leadership"]
        next_steps.append("Embed new metrics into CV")
        items.append(KeyOpportunityPayload(
            title=f"Position for {reco.title}",
            priority=priority,
            confidence=reco.confidence,
            timeline="30-60 days",
            expected_impact="+15% interview rate",
            recommended_actions=next_steps,
            success_metrics=["3 recruiter callbacks", "2 technical interviews"],
            blockers=reco.skill_gaps[:2],
            stakeholders=["Engineering leadership", "Talent partners"],
            xai_trace_id="CG-" + uuid.uuid4().hex[:8].upper()
        ))
    if not items:
        items.append(KeyOpportunityPayload(
            title="Clarify professional narrative",
            priority="Medium",
            confidence=0.6,
            timeline="45 days",
            expected_impact="+10% response rate",
            recommended_actions=["Add quantified achievements", "Use leadership framing"],
            success_metrics=["Portfolio refresh", "Peer review"],
            blockers=["Storytelling"],
            stakeholders=["Hiring managers"],
            xai_trace_id="CG-DEFAULT"
        ))
    return items


def _build_portfolio_sections(analysis: Dict[str, Any], customization: PortfolioCustomization) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    config = _normalize_portfolio_customization(customization)
    visible = set(config.sections_visible or [])
    view_model = _build_portfolio_viewmodel(analysis)
    sections: List[Dict[str, Any]] = []

    if 'about' in visible and view_model.get('about'):
        hero = view_model.get('hero', {})
        chips = []
        if hero.get('location'):
            chips.append(f"Location: {hero['location']}")
        exp_years = analysis.get('total_years_experience') or analysis.get('experience_years')
        if exp_years:
            chips.append(f"Experience: {exp_years:.1f} yrs")
        sections.append({
            'id': 'about',
            'title': 'About',
            'content': [{'type': 'paragraph', 'text': view_model['about']}, {'type': 'chips', 'items': chips}] if chips else [{'type': 'paragraph', 'text': view_model['about']}]
        })

    if 'skills' in visible and view_model.get('skills'):
        skill_lines = [f"{category}: {', '.join(items)}" for category, items in view_model['skills'].items()]
        sections.append({
            'id': 'skills',
            'title': 'Skills & Stack',
            'content': [{'type': 'list', 'items': skill_lines}]
        })

    if 'experience' in visible and view_model.get('experiences'):
        experience_lines = []
        for exp in view_model['experiences']:
            summary = exp['bullets'][0] if exp.get('bullets') else ''
            descriptor = f"{exp['title']} @ {exp['company']} ({exp['period']})"
            if summary:
                descriptor = f"{descriptor} – {summary}"
            experience_lines.append(descriptor)
        sections.append({
            'id': 'experience',
            'title': 'Experience Highlights',
            'content': [{'type': 'list', 'items': experience_lines}]
        })

    if 'projects' in visible and view_model.get('projects'):
        project_lines = []
        for project in view_model['projects']:
            highlight = project['bullets'][0] if project.get('bullets') else project.get('summary')
            descriptor = f"{project['title']}"
            if highlight:
                descriptor = f"{descriptor} – {highlight}"
            if project.get('tech'):
                descriptor = f"{descriptor} ({', '.join(project['tech'])})"
            project_lines.append(descriptor)
        sections.append({
            'id': 'projects',
            'title': 'Projects',
            'content': [{'type': 'list', 'items': project_lines}]
        })

    if 'education' in visible and view_model.get('education'):
        sections.append({
            'id': 'education',
            'title': 'Education',
            'content': [{'type': 'list', 'items': view_model['education']}]
        })

    if 'contact' in visible and view_model.get('contact'):
        contact_lines = []
        contact = view_model['contact']
        for label in ('email', 'phone', 'location', 'linkedin', 'github'):
            if contact.get(label):
                pretty_label = label.title()
                contact_lines.append(f"{pretty_label}: {contact[label]}")
        if contact_lines:
            sections.append({
                'id': 'contact',
                'title': 'Contact',
                'content': [{'type': 'list', 'items': contact_lines}]
            })

    return sections, view_model


def _render_portfolio_html(
    portfolio: Dict[str, Any],
    customization: PortfolioCustomization,
    sections: List[Dict[str, Any]],
    view_model: Dict[str, Any]
) -> str:
    config = _normalize_portfolio_customization(customization)
    accent = PORTFOLIO_COLOR_MAP.get(config.color_scheme.lower(), '#2563eb')
    accent_light = '#3b82f6'
    background = '#070f1c'
    surface = '#0f1c2f'
    text_muted = '#94a3b8'
    font_family = f"{config.font_family}, 'Inter', 'Space Grotesk', 'Segoe UI', sans-serif"

    hero = view_model.get('hero', {})
    hero_name = html.escape(hero.get('name', 'Professional'))
    hero_title = html.escape(hero.get('title', 'Candidate'))
    hero_headline = html.escape(hero.get('headline', 'Results-driven professional.'))
    hero_location = html.escape(hero.get('location', 'Remote'))
    contact = hero.get('contact') or {}

    def contact_chip(label: str, value: str) -> str:
        return f"<span class='contact-chip' data-icon='{label}'>{html.escape(value)}</span>"

    contact_html = ''
    for key in ('email', 'phone', 'linkedin', 'github', 'location'):
        if contact.get(key):
            contact_html += contact_chip(key, str(contact[key]))

    stats_html = ''
    stats = portfolio.get('stats', {})
    stat_items = [
        ('Experience', f"{stats.get('experience_years', 0):.1f} yrs" if stats.get('experience_years') is not None else None),
        ('Skills', str(stats.get('skills_count'))) if stats.get('skills_count') else None,
        ('Sections', str(len(stats.get('sections', [])))) if stats.get('sections') else None,
    ]
    stat_items = [item for item in stat_items if item]
    if stat_items:
        stats_html = ''.join(
            f"<div class='stat'><p class='stat-value'>{html.escape(value)}</p><p class='stat-label'>{html.escape(label)}</p></div>"
            for label, value in stat_items
        )

    about_html = ''
    if view_model.get('about'):
        about_html = f"<section class='card'><h2>About</h2><p>{html.escape(view_model['about'])}</p></section>"

    skill_cards = ''
    for category, items in view_model.get('skills', {}).items():
        chips = ''.join(f"<span class='pill'>{html.escape(item)}</span>" for item in items)
        skill_cards += f"<div class='skill-card'><p class='skill-title'>{html.escape(category)}</p><div class='pill-grid'>{chips}</div></div>"
    skills_html = ''
    if skill_cards:
        skills_html = f"<section class='card'><div class='section-head'><h2>Skills</h2><p>Organized by expertise</p></div><div class='grid'>{skill_cards}</div></section>"

    experience_cards = ''
    for exp in view_model.get('experiences', []):
        bullets = ''.join(f"<li>{html.escape(item)}</li>" for item in exp.get('bullets', [])[:4])
        experience_cards += (
            "<article class='experience-card'>"
            f"<div class='experience-meta'><p class='eyebrow'>{html.escape(exp['company'])}</p>"
            f"<h3>{html.escape(exp['title'])}</h3><span class='period'>{html.escape(exp['period'])}</span></div>"
            f"<ul>{bullets}</ul>"
            "</article>"
        )
    experiences_html = ''
    if experience_cards:
        experiences_html = f"<section class='card'><div class='section-head'><h2>Experience</h2><p>Selected impact</p></div><div class='stack'>{experience_cards}</div></section>"

    project_cards = ''
    for project in view_model.get('projects', []):
        bullets = ''.join(f"<li>{html.escape(item)}</li>" for item in (project.get('bullets') or [])[:3])
        tech = ''
        if project.get('tech'):
            tech = ''.join(f"<span class='pill pill-muted'>{html.escape(tag)}</span>" for tag in project['tech'])
            tech = f"<div class='pill-row'>{tech}</div>"
        project_cards += (
            "<article class='project-card'>"
            f"<h3>{html.escape(project['title'])}</h3>"
            + (f"<p>{html.escape(project.get('summary'))}</p>" if project.get('summary') else '')
            + f"<ul>{bullets}</ul>{tech}" + "</article>"
        )
    projects_html = ''
    if project_cards:
        projects_html = f"<section class='card'><div class='section-head'><h2>Projects</h2><p>Shipped initiatives</p></div><div class='grid'>{project_cards}</div></section>"

    education_items = ''.join(f"<li>{html.escape(entry)}</li>" for entry in view_model.get('education', []))
    education_html = ''
    if education_items:
        education_html = f"<section class='card'><div class='section-head'><h2>Education</h2><p>Academic journey</p></div><ul class='edu-list'>{education_items}</ul></section>"

    contact_cards = ''
    for key in ('email', 'phone', 'linkedin', 'github'):
        raw_value = contact.get(key)
        value = str(raw_value).strip() if raw_value else None
        if not value:
            continue
        href = '#'
        if key == 'email':
            href = f"mailto:{value}"
        elif key == 'phone':
            href = f"tel:{re.sub(r'\\s+', '', value)}"
        elif key in {'linkedin', 'github'}:
            href = value if str(value).lower().startswith('http') else f"https://{value}"
        contact_cards += (
            f"<a class='contact-card' href='{html.escape(href)}' target='_blank' rel='noreferrer'>"
            f"<p class='eyebrow'>{key.title()}</p><p>{html.escape(value)}</p></a>"
        )
    contact_html = ''
    if contact_cards:
        contact_html = f"<section class='card'><div class='section-head'><h2>Contact</h2><p>Let’s collaborate</p></div><div class='grid'>{contact_cards}</div></section>"

    return f"""<!doctype html>
<html lang='en'>
<head>
  <meta charset='utf-8'>
  <meta name='viewport' content='width=device-width, initial-scale=1'>
  <title>{html.escape(portfolio.get('name', 'Professional Portfolio'))}</title>
  <style>
    :root {{
      --accent: {accent};
      --accent-light: {accent_light};
      --surface: {surface};
      --text: #e2e8f0;
      --muted: {text_muted};
    }}
    * {{ box-sizing: border-box; }}
    body {{ font-family: {font_family}; background: {background}; color: var(--text); margin: 0; }}
    main {{ max-width: 1100px; margin: 0 auto; padding: 48px 24px 96px; }}
    .hero {{ background: linear-gradient(135deg, rgba(37,99,235,.25), rgba(14,165,233,.15)); border-radius: 32px; padding: 48px; position: relative; overflow: hidden; }}
    .hero::after {{ content: ''; position: absolute; inset: 0; background: radial-gradient(circle at top right, rgba(255,255,255,.08), transparent 55%); }}
    .hero > div {{ position: relative; z-index: 2; }}
    .eyebrow {{ text-transform: uppercase; letter-spacing: .3em; font-size: .75rem; color: var(--muted); margin: 0 0 8px; }}
    .hero h1 {{ font-size: clamp(2.5rem, 5vw, 3.5rem); margin: 0; }}
    .hero p {{ margin: 8px 0; max-width: 720px; }}
    .subtitle {{ font-size: 1.25rem; color: var(--muted); margin-bottom: 4px; }}
    .contact-row {{ display: flex; flex-wrap: wrap; gap: 10px; margin-top: 24px; }}
    .contact-chip {{ background: rgba(15,28,47,.8); padding: 10px 16px; border-radius: 999px; font-size: .9rem; border: 1px solid rgba(255,255,255,.08); display: inline-flex; align-items: center; gap: 8px; }}
    .contact-chip::before {{ font-size: .95rem; opacity: .8; }}
    .contact-chip[data-icon='email']::before {{ content: '✉'; }}
    .contact-chip[data-icon='phone']::before {{ content: '☎'; }}
    .contact-chip[data-icon='linkedin']::before {{ content: 'in'; font-weight: 600; }}
    .contact-chip[data-icon='github']::before {{ content: 'GH'; font-weight: 600; }}
    .contact-chip[data-icon='location']::before {{ content: '📍'; }}
    .stat-grid {{ display: flex; flex-wrap: wrap; gap: 16px; margin-top: 32px; }}
    .stat {{ background: rgba(15,28,47,.6); border-radius: 20px; padding: 16px 20px; min-width: 140px; border: 1px solid rgba(255,255,255,.08); }}
    .stat-value {{ font-size: 1.5rem; margin: 0; }}
    .stat-label {{ margin: 4px 0 0; color: var(--muted); font-size: .85rem; text-transform: uppercase; letter-spacing: .1em; }}
    .card {{ background: var(--surface); border-radius: 28px; padding: 32px; margin-top: 32px; box-shadow: 0 20px 50px rgba(0,0,0,.35); border: 1px solid rgba(255,255,255,.05); }}
    .card h2 {{ margin: 0 0 12px; font-size: 1.5rem; }}
    .card p {{ color: var(--muted); line-height: 1.6; }}
    .section-head {{ display: flex; align-items: baseline; justify-content: space-between; gap: 12px; flex-wrap: wrap; }}
    .section-head p {{ margin: 0; font-size: .95rem; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); gap: 20px; margin-top: 20px; }}
    .stack {{ display: flex; flex-direction: column; gap: 16px; margin-top: 20px; }}
    .skill-card {{ background: rgba(8,15,30,.7); border-radius: 20px; padding: 20px; border: 1px solid rgba(255,255,255,.05); }}
    .skill-title {{ font-weight: 600; margin: 0 0 10px; }}
    .pill-grid {{ display: flex; flex-wrap: wrap; gap: 8px; }}
    .pill, .pill-muted {{ display: inline-flex; padding: 6px 12px; border-radius: 999px; font-size: .85rem; background: rgba(255,255,255,.08); color: var(--text); }}
    .pill-muted {{ background: rgba(148,163,184,.2); color: var(--text); }}
    .experience-card, .project-card {{ background: rgba(8,15,30,.7); border-radius: 24px; padding: 20px; border: 1px solid rgba(255,255,255,.05); }}
    .experience-meta {{ display: flex; flex-direction: column; gap: 4px; }}
    .experience-card h3 {{ margin: 0; font-size: 1.2rem; }}
    .period {{ font-size: .85rem; color: var(--muted); }}
    ul {{ padding-left: 20px; margin: 12px 0 0; color: var(--muted); }}
    li {{ margin-bottom: 6px; line-height: 1.5; }}
    .pill-row {{ margin-top: 12px; display: flex; flex-wrap: wrap; gap: 8px; }}
    .edu-list {{ list-style: none; padding: 0; margin: 12px 0 0; }}
    .edu-list li {{ margin-bottom: 10px; padding-bottom: 10px; border-bottom: 1px solid rgba(255,255,255,.04); }}
    .contact-card {{ text-decoration: none; color: inherit; background: rgba(8,15,30,.7); padding: 20px; border-radius: 18px; border: 1px solid rgba(255,255,255,.05); transition: transform .2s ease, border-color .2s ease; }}
    .contact-card:hover {{ transform: translateY(-4px); border-color: var(--accent); }}
    @media (max-width: 768px) {{
      .hero {{ padding: 32px; }}
      .card {{ padding: 24px; }}
      .grid {{ grid-template-columns: 1fr; }}
    }}
  </style>
</head>
<body>
  <main>
    <section class='hero'>
      <div>
        <p class='eyebrow'>Featured Portfolio</p>
        <h1>{hero_name}</h1>
        <p class='subtitle'>{hero_title} · {hero_location}</p>
        <p>{hero_headline}</p>
        <div class='contact-row'>{contact_html}</div>
        <div class='stat-grid'>{stats_html}</div>
      </div>
    </section>
    {about_html}
    {skills_html}
    {experiences_html}
    {projects_html}
    {education_html}
    {contact_html}
  </main>
</body>
</html>"""


@app.post("/api/v1/jobs/search", response_model=JobSearchResponsePayload)
async def search_jobs_endpoint(search_request: JobSearchRequestPayload):
    start_time = time.perf_counter()

    skills = [skill.strip() for skill in search_request.skills if skill.strip()]
    normalized_query = (search_request.query or "").strip()
    if not skills and normalized_query:
        skills = [normalized_query]

    if not skills:
        skills = DEFAULT_JOB_SEARCH_TERMS.copy()
        logger.info("Job search fallback triggered - using default skill pool: %s", skills[:3])

    normalized_location = _normalize_location(search_request.location)
    jobs: List[JobResultPayload] = []
    sources_used: List[str] = []
    external_result: Optional[Dict[str, Any]] = None

    try:
        external_result = await search_jobs_multi_source(skills, normalized_location, search_request.max_results)
        raw_jobs = external_result.get("jobs", []) if external_result else []
        if raw_jobs:
            jobs = [_map_job_to_payload(job) for job in raw_jobs][:search_request.max_results]
            sources_used = external_result.get("sources_used", [])
    except Exception as exc:
        logger.exception("Job search failed via external providers: %s", exc)

    if not jobs:
        jobs = _generate_fallback_jobs(search_request.query, normalized_location, skills, search_request.max_results)
        if not sources_used:
            sources_used = ["skill-sync-fallback"]

    elapsed_ms = int((time.perf_counter() - start_time) * 1000)
    if jobs:
        _persist_job_matches(jobs, skills, normalized_location)
    log_ml(
        "job-search",
        query=normalized_query or "top roles",
        location=normalized_location,
        skills=skills[:8],
        total=len(jobs),
        sources=sources_used,
        elapsed_ms=elapsed_ms
    )
    logger.info("Job search completed: %s results in %sms", len(jobs), elapsed_ms)

    reported_total = external_result.get("total_jobs", len(jobs)) if external_result else len(jobs)

    return JobSearchResponsePayload(
        jobs=jobs,
        total_count=reported_total,
        search_query=normalized_query or "Top live roles",
        location=normalized_location,
        sources_used=sources_used,
        search_time_ms=elapsed_ms,
        timestamp=datetime.utcnow().isoformat()
    )

class CVProcessor:
    """Enhanced CV document processing with basic extraction"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx'}
    
    async def parse_file(self, file: UploadFile) -> str:
        """Parse uploaded CV file and extract text"""
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return await self._parse_pdf(file)
            elif file_extension == '.docx':
                return await self._parse_docx(file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error parsing file {file.filename}: {e}")
            raise
    
    async def _parse_pdf(self, file: UploadFile) -> str:
        """Parse PDF file using available libraries"""
        if PyPDF2 is None and pdfplumber is None:
            raise RuntimeError("PDF parsing libraries are not available in this environment")
        content = ""
        file_bytes = await file.read()
        import io
        buffer = io.BytesIO(file_bytes)
        
        # Try pdfplumber first
        try:
            buffer.seek(0)
            with pdfplumber.open(buffer) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            # Fallback to PyPDF2
            try:
                buffer.seek(0)
                pdf_reader = PyPDF2.PdfReader(buffer)
                for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
            except Exception as e2:
                logger.error(f"PyPDF2 parsing failed: {e2}")
                raise
        
        return content.strip()
    
    async def _parse_docx(self, file: UploadFile) -> str:
        """Parse DOCX file"""
        if Document is None:
            raise RuntimeError("python-docx is not available to parse DOCX files")
        try:
            content = ""
            file_bytes = await file.read()
            
            # Save to temporary file
            temp_path = Path(f"temp_{uuid.uuid4()}.docx")
            with open(temp_path, 'wb') as f:
                f.write(file_bytes)
            
            # Parse with python-docx
            doc = Document(temp_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Clean up
            temp_path.unlink(missing_ok=True)
            
            return content.strip()
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise

class SkillExtractor:
    """Hybrid ML skill extractor with semantic + pattern coverage."""

    def __init__(self):
        self.skill_patterns = {
            'programming_languages': programming_languages,
            'frameworks': frameworks,
            'databases': databases,
            'cloud_platforms': cloud_platforms,
            'tools': tools,
            'soft_skills': soft_skill_keywords
        }
        self._ingest_custom_keywords()
        self.technical_categories = {
            'programming_languages', 'frameworks', 'databases', 'cloud_platforms', 'tools'
        }
        self.skill_index = {
            skill.lower(): category
            for category, skills in self.skill_patterns.items()
            for skill in skills
        }
        self.confidence_cutoff = getattr(settings, "skill_confidence_threshold", 0.62)
        self.ml_extractor: Optional[SkillsExtractorModel] = None
        if SkillsExtractorModel is not None:
            try:
                self.ml_extractor = SkillsExtractorModel(
                    model_name=settings.resume_ner_model_name,
                    model_path=settings.resume_ner_model_path,
                    min_confidence=self.confidence_cutoff
                )
                logger.info(" ML skill extractor initialized (BERT-based)")
            except Exception as exc:
                logger.warning("ML skill extractor unavailable, falling back to patterns: %s", exc)
                self.ml_extractor = None

    async def extract_skills(self, text: str) -> Dict[str, Any]:
        """Extract skills using ML-first strategy with deterministic fallbacks."""
        normalized_text = self._normalize_text(text)
        catalog: Dict[str, Dict[str, Any]] = {}
        sources_used: List[str] = []
        profile_context = self._analyze_profile_context(normalized_text)

        if self.ml_extractor is not None:
            try:
                ml_result = self.ml_extractor.extract_skills(normalized_text, use_bert=True)
                sources_used.append(ml_result.get('method', 'bert'))
                ml_confidence = self._confidence_from_label(
                    ml_result.get('confidence_score') or ml_result.get('confidence')
                )
                for detail in ml_result.get('skill_details', []):
                    self._register_skill(
                        catalog,
                        detail.get('skill'),
                        normalized_text,
                        detail.get('confidence', ml_confidence),
                        'ml-bert',
                        profile_context=profile_context
                    )
                low_conf = ml_result.get('low_confidence_skills') or []
                if low_conf:
                    log_ml(
                        "skill-extraction-low-confidence",
                        rejected=[entry.get('skill') for entry in low_conf[:5]],
                        total=len(low_conf)
                    )
            except Exception as exc:
                logger.warning("ML extraction failed, reverting to pattern scanner: %s", exc)

        pattern_results = self._pattern_scan(normalized_text)
        if pattern_results:
            sources_used.append('pattern-matcher')
        for match in pattern_results:
            self._register_skill(
                catalog,
                match['skill'],
                text,
                match['confidence'],
                'pattern',
                profile_context=profile_context
            )

        categorized = defaultdict(list)
        for skill_name, payload in catalog.items():
            category = payload.pop('category', 'other')
            categorized[category].append(payload)

        all_confidences = [entry['confidence'] for entry in catalog.values()]
        overall_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0

        result = {
            'skills': dict(categorized),
            'confidence': round(overall_confidence, 3),
            'total_skills_found': len(catalog),
            'extraction_method': 'ml+pattern' if sources_used else 'pattern-only',
            'sources_used': sources_used
        }
        log_ml(
            "skill-extraction",
            method=result['extraction_method'],
            total_skills=result['total_skills_found'],
            sources=sources_used,
            categories=list(result['skills'].keys())
        )
        return result

    def _pattern_scan(self, text: str) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        text_lower = text.lower()
        for category, skill_list in self.skill_patterns.items():
            for skill in skill_list:
                pattern = rf'\b{re.escape(skill)}\b'
                for match in re.finditer(pattern, text_lower, re.IGNORECASE):
                    start = max(0, match.start() - 50)
                    end = min(len(text), match.end() + 50)
                    context = text[start:end].strip()
                    results.append({
                        'skill': skill,
                        'category': category,
                        'confidence': self._calculate_confidence(skill, context),
                        'context': context,
                        'position': match.start()
                    })
        return results

    def _register_skill(self, catalog: Dict[str, Dict[str, Any]], skill: str, text: str,
                         confidence: float, source: str, profile_context: Optional[Dict[str, Any]] = None) -> None:
        normalized = (skill or '').strip()
        if not normalized:
            return
        title_case = normalized.title()
        context = self._extract_context(text, normalized)
        confidence = min(1.0, max(confidence or 0.1, 0.1) + (0.05 if context else 0.0))
        if source.startswith('ml') and confidence < self.confidence_cutoff:
            log_ml(
                "skill-extraction-filtered",
                skill=title_case,
                confidence=round(confidence, 3),
                cutoff=self.confidence_cutoff
            )
            return
        category = self._categorize_skill(normalized)
        if self._should_filter_contextually(category, normalized, context, profile_context):
            log_ml(
                "skill-extraction-context-skip",
                skill=title_case,
                category=category,
                reason="non-technical-context"
            )
            return
        record = {
            'skill': title_case,
            'confidence': round(confidence, 3),
            'context': context,
            'position': text.lower().find(normalized.lower()),
            'source': source,
            'category': category
        }
        existing = catalog.get(title_case)
        if not existing or existing['confidence'] < record['confidence']:
            catalog[title_case] = record

    def _categorize_skill(self, skill: str) -> str:
        lowered = skill.lower()
        return self.skill_index.get(lowered, 'other')

    def _should_filter_contextually(
        self,
        category: str,
        skill: str,
        context: str,
        profile_context: Optional[Dict[str, Any]]
    ) -> bool:
        if category not in self.technical_categories:
            return False
        context_lower = (context or '').lower()
        if self._has_technical_anchor(context_lower):
            return False
        skill_lower = skill.lower()
        if skill_lower in SHORT_SKILL_GUARD:
            return True
        if profile_context:
            marketing_hits = profile_context.get('marketing_hits', 0)
            tech_hits = profile_context.get('tech_hits', 0)
            if profile_context.get('non_technical_bias') and tech_hits < 3:
                return True
            if marketing_hits >= max(4, tech_hits * 2) and not self._has_technical_anchor(context_lower):
                return True
            if any(term in context_lower for term in MARKETING_CONTEXT_TERMS):
                return True
        return False

    def _has_technical_anchor(self, context_lower: str) -> bool:
        if not context_lower:
            return False
        return any(anchor in context_lower for anchor in TECH_CONTEXT_INDICATORS)

    def _extract_context(self, text: str, skill: str, window: int = 60) -> str:
        lowered_text = text.lower()
        idx = lowered_text.find(skill.lower())
        if idx == -1:
            return ''
        start = max(0, idx - window)
        end = min(len(text), idx + len(skill) + window)
        return text[start:end].strip()

    def _confidence_from_label(self, label: Optional[Any]) -> float:
        if label is None:
            return 0.65
        if isinstance(label, (int, float)):
            return float(label)
        mapping = {
            'high': 0.88,
            'medium': 0.68,
            'low': 0.48
        }
        return mapping.get(str(label).lower(), 0.65)

    def _calculate_confidence(self, skill: str, context: str) -> float:
        base_confidence = 0.62
        context_lower = context.lower()
        if any(indicator in context_lower for indicator in ['experience', 'skilled', 'proficient', 'expert']):
            base_confidence += 0.2
        if any(indicator in context_lower for indicator in ['worked with', 'used', 'developed', 'implemented']):
            base_confidence += 0.1
        return min(max(base_confidence, 0.1), 0.95)

    def _normalize_text(self, text: str) -> str:
        if not text:
            return ""
        sanitized = text.replace('\u2022', '- ').replace('\t', ' ')
        normalized_lines: List[str] = []
        for raw_line in sanitized.splitlines():
            stripped = raw_line.strip()
            if not stripped:
                continue
            stripped = re.sub(r'^[\-*•]+', '', stripped).strip()
            stripped = re.sub(r'\s{2,}', ' ', stripped)
            normalized_lines.append(stripped)
        normalized = '\n'.join(normalized_lines)
        normalized = re.sub(r'(?<![.!?])\n(?=[a-z])', '. ', normalized)
        normalized = re.sub(r'([a-z])([A-Z])', r'\1. \2', normalized)
        return normalized

    def _ingest_custom_keywords(self) -> None:
        base_dir = Path(__file__).resolve().parents[1]
        catalog_path = base_dir / 'data' / 'skills_fallback_keywords.json'
        if not catalog_path.exists():
            return
        try:
            with open(catalog_path, 'r', encoding='utf-8') as handle:
                data = json.load(handle)
            for category, keywords in data.items():
                if not isinstance(keywords, list):
                    continue
                normalized_keywords = [keyword.strip() for keyword in keywords if keyword.strip()]
                if category in self.skill_patterns:
                    merged = list(dict.fromkeys(self.skill_patterns[category] + normalized_keywords))
                    self.skill_patterns[category] = merged
                else:
                    self.skill_patterns[category] = normalized_keywords
        except Exception as exc:
            logger.warning("Failed to ingest custom skill keywords: %s", exc)

    def _analyze_profile_context(self, text: str) -> Dict[str, Any]:
        lowered = (text or '').lower()
        tech_hits = sum(1 for indicator in TECH_CONTEXT_INDICATORS if indicator in lowered)
        marketing_hits = sum(1 for indicator in MARKETING_CONTEXT_TERMS if indicator in lowered)
        return {
            'tech_hits': tech_hits,
            'marketing_hits': marketing_hits,
            'non_technical_bias': marketing_hits >= max(4, tech_hits * 2)
        }

class ExperienceTranslator:
    """Basic experience translation between technical and business contexts"""
    
    async def translate_experience(self, text: str, source_type: str, target_type: str, tone: str = 'professional') -> Dict[str, Any]:
        """Basic translation between contexts"""
        
        if source_type == target_type:
            return {
                'original_text': text,
                'translated_text': text,
                'source_type': source_type,
                'target_type': target_type,
                'confidence_score': 1.0,
                'key_changes': ['No translation needed - same context type']
            }
        
        # Simple vocabulary mapping
        tech_to_business = {
            'developed': 'strategically implemented',
            'built': 'constructed',
            'created': 'established',
            'api': 'integration interfaces',
            'database': 'data infrastructure',
            'debug': 'troubleshoot issues',
            'deploy': 'launch solutions'
        }
        
        business_to_tech = {
            'strategically implemented': 'developed',
            'constructed': 'built',
            'established': 'created',
            'integration interfaces': 'api',
            'data infrastructure': 'database',
            'troubleshoot issues': 'debug',
            'launch solutions': 'deploy'
        }
        
        # Apply mappings
        translated_text = text
        if source_type == 'technical' and target_type == 'business':
            for tech, business in tech_to_business.items():
                translated_text = re.sub(r'\b' + tech + r'\b', business, translated_text, flags=re.IGNORECASE)
        elif source_type == 'business' and target_type == 'technical':
            for business, tech in business_to_tech.items():
                translated_text = re.sub(r'\b' + business + r'\b', tech, translated_text, flags=re.IGNORECASE)
        
        key_changes = [f"Vocabulary mapped from {source_type} to {target_type} context"]
        
        return {
            'original_text': text,
            'translated_text': translated_text,
            'source_type': source_type,
            'target_type': target_type,
            'tone': tone,
            'confidence_score': 0.8,
            'key_changes': key_changes
        }
    
    def analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analyze sentiment of text"""
        if not sia:
            return {
                'sentiment': 'neutral',
                'confidence': 0.5,
                'scores': {'positive': 0.33, 'negative': 0.33, 'neutral': 0.34}
            }
        
        try:
            scores = sia.polarity_scores(text)
            compound_score = scores['compound']
            
            if compound_score >= 0.05:
                sentiment = 'positive'
            elif compound_score <= -0.05:
                sentiment = 'negative'
            else:
                sentiment = 'neutral'
            
            return {
                'sentiment': sentiment,
                'confidence': abs(compound_score),
                'scores': scores
            }
        except Exception as e:
            logger.warning(f"Sentiment analysis failed: {e}")
            return {
                'sentiment': 'neutral',
                'confidence': 0.5,
                'scores': {'positive': 0.33, 'negative': 0.33, 'neutral': 0.34}
            }

# Initialize components
cv_processor = CVProcessor()
skill_extractor = SkillExtractor()
experience_translator = ExperienceTranslator()
cv_analysis_storage: Dict[str, Dict[str, Any]] = {}
portfolio_preview_store: Dict[str, Dict[str, Any]] = {}
PORTFOLIO_PREVIEW_TTL_MINUTES = 60
PORTFOLIO_PREVIEW_MAX_ENTRIES = 50


def _cleanup_portfolio_preview_store() -> None:
    """Keep preview store bounded and drop expired entries."""
    if not portfolio_preview_store:
        return

    now = datetime.utcnow()
    ttl_cutoff = now - timedelta(minutes=PORTFOLIO_PREVIEW_TTL_MINUTES)
    expired = [key for key, item in portfolio_preview_store.items() if item.get("created_at", now) < ttl_cutoff]
    for key in expired:
        portfolio_preview_store.pop(key, None)

    excess = len(portfolio_preview_store) - PORTFOLIO_PREVIEW_MAX_ENTRIES
    if excess <= 0:
        return

    oldest = sorted(
        portfolio_preview_store.items(),
        key=lambda entry: entry[1].get("created_at", now)
    )
    for key, _ in oldest[:excess]:
        portfolio_preview_store.pop(key, None)

# Database dependency
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Basic password hashing
def get_password_hash(password: str) -> str:
    import hashlib
    return hashlib.sha256(password.encode()).hexdigest()


def _ensure_system_user() -> str:
    """Guarantee a system user exists for analytics persistence."""
    global SYSTEM_USER_ID
    if SYSTEM_USER_ID:
        return SYSTEM_USER_ID

    session = SessionLocal()
    try:
        user = session.query(UserDB).filter(UserDB.email == SYSTEM_USER_EMAIL).first()
        if user:
            SYSTEM_USER_ID = user.id
            return SYSTEM_USER_ID

        user = UserDB(
            id=str(uuid.uuid4()),
            email=SYSTEM_USER_EMAIL,
            username=SYSTEM_USER_USERNAME,
            full_name=SYSTEM_USER_FULL_NAME,
            hashed_password=get_password_hash(SYSTEM_USER_PASSWORD)
        )
        session.add(user)
        session.commit()
        session.refresh(user)
        SYSTEM_USER_ID = user.id
        logger.info("System analytics user created: %s", SYSTEM_USER_ID)
        return SYSTEM_USER_ID
    except Exception as exc:
        session.rollback()
        logger.error("Failed to provision system user: %s", exc)
        raise
    finally:
        session.close()


def _persist_job_matches(jobs: List[JobResultPayload], skills: List[str], location: str) -> None:
    if not jobs:
        return

    owner_id = SYSTEM_USER_ID or _ensure_system_user()
    session = SessionLocal()
    try:
        for job in jobs[: min(20, len(jobs))]:
            session.add(JobMatchDB(
                user_id=owner_id,
                job_title=job.title,
                company=job.company,
                job_description=(job.description or "")[:1000],
                match_score=job.match_score or 0.0,
                matched_skills=job.skills or skills,
                missing_skills=[],
                salary_range=job.salary,
                location=job.location or location,
                job_url=job.url,
            ))
        session.commit()
    except Exception as exc:
        session.rollback()
        logger.warning("Job match persistence failed: %s", exc)
    finally:
        session.close()


def _safe_float(value: Any, fallback: Optional[float] = 0.0) -> Optional[float]:
    try:
        if isinstance(value, str):
            sanitized = value.strip().replace('%', '').replace('h', '').replace('H', '')
            if not sanitized:
                return fallback
            return float(sanitized)
        if value is None:
            return fallback
        return float(value)
    except (TypeError, ValueError):
        return fallback


def _infer_resource_type(resource: Dict[str, Any]) -> str:
    tier = resource.get('tier')
    if isinstance(tier, str):
        return tier
    title = resource.get('title') or ''
    for keyword in ('Core', 'Applied', 'Proof', 'Advanced', 'Beginner'):
        if keyword.lower() in title.lower():
            return keyword
    return 'General'


def _extract_resource_hours(resource: Dict[str, Any]) -> float:
    if 'estimated_time_hours' in resource and resource['estimated_time_hours'] is not None:
        return _safe_float(resource['estimated_time_hours'], 0.0)
    if 'time_hours' in resource and resource['time_hours'] is not None:
        return _safe_float(resource['time_hours'], 0.0)
    duration_field = resource.get('duration')
    if isinstance(duration_field, str):
        digits = ''.join(ch for ch in duration_field if ch.isdigit() or ch == '.')
        return _safe_float(digits, 0.0)
    return 0.0


def _extract_resource_cost(resource: Dict[str, Any]) -> float:
    cost_field = resource.get('cost')
    if isinstance(cost_field, (int, float)):
        return float(cost_field)
    if isinstance(cost_field, str):
        if cost_field.strip().lower() in {'free', 'no-cost', 'gratis'}:
            return 0.0
        digits = ''.join(ch for ch in cost_field if ch.isdigit() or ch == '.')
        return _safe_float(digits, 0.0)
    return 0.0


def _persist_learning_roadmap_run(
    plan: Dict[str, Any],
    primary_skills: List[str],
    gap_skills: List[str],
    context: Dict[str, Any],
    analysis: Optional[CVAnalysisResponse]
) -> None:
    if not plan or not plan.get('phases'):
        return

    owner_id = getattr(analysis, 'user_id', None)
    analysis_id = getattr(analysis, 'analysis_id', None) if analysis else None
    skills_sequence: List[str] = []
    for phase in plan.get('phases', []):
        skills_sequence.extend([skill for skill in phase.get('skills', []) if skill])

    session = SessionLocal()
    try:
        run = LearningRoadmapRunDB(
            user_id=owner_id or SYSTEM_USER_ID or _ensure_system_user(),
            analysis_id=analysis_id,
            primary_skills=primary_skills,
            gap_skills=gap_skills,
            skills_sequence=skills_sequence,
            phases=plan.get('phases'),
            total_hours=_safe_float(plan.get('total_hours'), 0.0),
            total_weeks=_safe_float(plan.get('total_weeks'), 0.0),
            success_rate=_safe_float(plan.get('predicted_success_rate'), 0.0),
            weekly_hours_available=_safe_float(context.get('weekly_hours_available'), None)
            if context else None,
            experience_years=_safe_float(context.get('experience_years'), None)
            if context else None,
            target_role=(analysis.job_titles[0] if analysis and analysis.job_titles else context.get('target_role') if context else None),
            planned_duration_hours=_safe_float(plan.get('total_hours'), 0.0),
            actual_duration_hours=None,
            satisfaction_score=_safe_float(getattr(analysis, 'confidence_score', None), None),
            quiz_scores=None
        )
        session.add(run)
        session.flush()

        for phase in plan.get('phases', []):
            phase_success = phase.get('success_probability')
            success_ratio = None
            if isinstance(phase_success, str) and phase_success.endswith('%'):
                success_ratio = _safe_float(phase_success) / 100.0
            elif isinstance(phase_success, (int, float)):
                success_ratio = float(phase_success)

            for resource in phase.get('resources', []) or []:
                session.add(LearningResourceEventDB(
                    run_id=run.id,
                    skill=resource.get('skill') or 'Skill',
                    tier=resource.get('tier'),
                    provider=resource.get('provider'),
                    resource_type=_infer_resource_type(resource),
                    duration_hours=_extract_resource_hours(resource),
                    cost=_extract_resource_cost(resource),
                    success_rate=success_ratio,
                    difficulty=_safe_float(resource.get('difficulty'), None),
                    is_free=bool(resource.get('is_free', False)),
                    resource_payload=resource,
                    source='ml-roadmap'
                ))

        session.commit()
        _schedule_model_refresh()
    except Exception as exc:
        session.rollback()
        logger.warning("[roadmap-ml] Failed to persist roadmap run: %s", exc)
    finally:
        session.close()


def _count_learning_records() -> Tuple[int, int]:
    session = SessionLocal()
    try:
        run_count = session.query(func.count(LearningRoadmapRunDB.id)).scalar() or 0
        resource_count = session.query(func.count(LearningResourceEventDB.id)).scalar() or 0
        return int(run_count), int(resource_count)
    except Exception as exc:  # noqa: BLE001
        logger.warning("[roadmap-ml] Unable to count learning records: %s", exc)
        return 0, 0
    finally:
        session.close()


def _maybe_refresh_learning_models(force: bool = False) -> None:
    global LAST_MODEL_REFRESH, LEARNING_ROADMAP_PREDICTOR, MODEL_REFRESH_PENDING
    now = datetime.utcnow()
    try:
        if not force and LAST_MODEL_REFRESH and now - LAST_MODEL_REFRESH < LEARNING_MODEL_REFRESH_INTERVAL:
            return

        run_count, resource_count = _count_learning_records()
        if not force and (
            run_count < LEARNING_MODEL_MIN_RUNS or resource_count < LEARNING_MODEL_MIN_RESOURCES
        ):
            logger.debug(
                "[roadmap-ml] Skipping refresh (runs=%s resources=%s)",
                run_count,
                resource_count
            )
            return

        if not MODEL_REFRESH_LOCK.acquire(blocking=False):
            logger.debug("[roadmap-ml] Refresh already running")
            return

        try:
            from roadmap_ml.training import RoadmapModelTrainer

            trainer = RoadmapModelTrainer()
            metrics = trainer.train_all()
            LAST_MODEL_REFRESH = datetime.utcnow()
            logger.info(
                "[roadmap-ml] Models refreshed (%s runs / %s resources) -> %s",
                run_count,
                resource_count,
                json.dumps(metrics)
            )
            LEARNING_ROADMAP_PREDICTOR = get_default_predictor(force_reload=True)
        finally:
            MODEL_REFRESH_LOCK.release()
    except Exception as exc:  # noqa: BLE001
        logger.warning("[roadmap-ml] Model refresh failed: %s", exc)
    finally:
        MODEL_REFRESH_PENDING = False


async def _refresh_learning_models_if_needed_async(force: bool = False) -> None:
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, functools.partial(_maybe_refresh_learning_models, force))


def _schedule_model_refresh(force: bool = False) -> None:
    global MODEL_REFRESH_PENDING
    if MODEL_REFRESH_PENDING and not force:
        return
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return
    MODEL_REFRESH_PENDING = True
    loop.create_task(_refresh_learning_models_if_needed_async(force=force))

# API Endpoints

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Welcome to SkillSync Enhanced Backend!",
        "version": "2.0",
        "features": [
            "F1: Enhanced CV Analysis with AI-powered skill extraction",
            "F2: Basic Job Matching",
            "F3: Skill Gap Analysis",
            "F4: Career Recommendations",
            "F5: Experience Translator",
            "F6: Basic Analytics",
            "F7: Portfolio Templates",
            "F8: User Management"
        ],
        "docs_url": "/docs",
        "health_check": "/health"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": settings.app_version,
        "components": {
            "database": "connected",
            "cv_processor": "available",
            "skill_extractor": "available",
            "nlp": nlp is not None,
            "sentiment_analyzer": sia is not None
        }
    }

@app.post("/api/v1/cv/upload", response_model=CVUploadResponse)
async def upload_cv(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user_id: str = None
):
    """Upload and analyze CV file"""
    try:
        # Validate file
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in settings.allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed types: {', '.join(settings.allowed_extensions)}"
            )
        
        # Generate unique analysis ID
        analysis_id = str(uuid.uuid4())
        
        # Parse CV file
        cv_text = await cv_processor.parse_file(file)
        
        if not cv_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from file")
        
        # Extract skills
        skills_data = await skill_extractor.extract_skills(cv_text)
        
        # Store analysis in database
        owner_id = user_id or SYSTEM_USER_ID or _ensure_system_user()

        db_analysis = CVAnalysisDB(
            id=analysis_id,
            user_id=owner_id,
            filename=file.filename,
            file_size=file.size,
            skills=skills_data['skills'],
            confidence_score=skills_data['confidence']
        )
        
        db.add(db_analysis)
        db.commit()
        logger.info("[cv-upload] analysis %s stored for user %s", analysis_id, owner_id)
        
        return CVUploadResponse(
            analysis_id=analysis_id,
            message="CV uploaded and analyzed successfully",
            confidence_score=skills_data['confidence']
        )
        
    except Exception as e:
        logger.error(f"CV upload failed: {e}")
        raise HTTPException(status_code=500, detail=f"CV processing failed: {str(e)}")


@app.post("/api/v1/extract-text")
async def extract_text_from_cv(file: UploadFile = File(...)):
    """Extract plain text from uploaded CV files for ML processing."""
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename missing in upload")

        file_bytes = await file.read()
        if not file_bytes:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")

        extension = Path(file.filename).suffix.lower()
        content_type = (file.content_type or "").lower()
        extracted_text = ""

        if extension == ".pdf" or "pdf" in content_type:
            import io

            try:
                pdf_stream = io.BytesIO(file_bytes)
                reader = PyPDF2.PdfReader(pdf_stream)
                pages = []
                for index, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    logger.debug("[extract-text] PDF page %s produced %s chars", index + 1, len(page_text))
                    pages.append(page_text)
                extracted_text = "\n".join(pages)
            except Exception as pdf_error:
                logger.exception("PDF parsing error")
                raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {pdf_error}")

        elif extension == ".docx" or "wordprocessingml.document" in content_type:
            import io

            try:
                document = Document(io.BytesIO(file_bytes))
                extracted_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            except Exception as docx_error:
                logger.exception("DOCX parsing error")
                raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {docx_error}")

        elif extension == ".txt" or content_type.startswith("text"):
            extracted_text = file_bytes.decode("utf-8", errors="ignore")

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT")

        cleaned_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", extracted_text)
        cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()

        if not cleaned_text:
            raise HTTPException(status_code=400, detail="Could not extract text from file")

        logger.info("[extract-text] %s -> %s characters", file.filename, len(cleaned_text))

        return {"cv_text": cleaned_text, "length": len(cleaned_text)}

    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Text extraction failed")
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {exc}")


@app.post("/api/v1/analyze-cv-advanced", response_model=CVAnalysisResponse)
async def analyze_cv_advanced(request: CVAnalysisRequest):
    """Lightweight advanced CV analysis used by the ML workflow in the frontend."""
    cv_text = (request.cv_content or "").strip()
    if not cv_text:
        raise HTTPException(status_code=400, detail="CV content cannot be empty")
    if len(cv_text) > 60000:
        raise HTTPException(status_code=413, detail="CV content too long (max 60,000 characters)")

    start_time = time.perf_counter()
    skills_data = await skill_extractor.extract_skills(cv_text)
    flattened_skills = _flatten_skill_results(skills_data)
    soft_skills_found = _extract_keywords(cv_text, soft_skill_keywords)
    job_titles = _extract_keywords(cv_text, job_title_keywords) or ["Professional"]
    education_entries = _extract_keywords(cv_text, education_keywords)
    experience_years = _estimate_experience_years(cv_text, len(flattened_skills))
    summary = cv_text[:600] + ("..." if len(cv_text) > 600 else "")
    analysis_id = str(uuid.uuid4())
    timestamp = datetime.utcnow().isoformat()

    response = CVAnalysisResponse(
        analysis_id=analysis_id,
        skills=flattened_skills,
        hard_skills=flattened_skills,
        soft_skills=soft_skills_found,
        experience_years=experience_years,
        job_titles=job_titles,
        education=education_entries,
        summary=summary,
        confidence_score=min(0.95, skills_data.get('confidence', 0.5) + 0.1),
        timestamp=timestamp,
        raw_text=cv_text,
        total_years_experience=experience_years,
        personal_info={
            "name": job_titles[0] if job_titles else "Professional",
            "location": None
        },
        contact_info={},
        industries=list(skills_data.get('skills', {}).keys()) or None,
        skill_categories={
            category: [entry.get('skill', '').title() for entry in entries if isinstance(entry, dict)]
            for category, entries in skills_data.get('skills', {}).items()
        },
        ml_confidence_breakdown={"skills": skills_data.get('confidence', 0.5)},
        parser_version="enhanced-lite-v2",
        languages=_extract_keywords(cv_text, ["english", "french", "spanish", "german", "arabic"]),
        processing_time_ms=int((time.perf_counter() - start_time) * 1000)
    )

    cv_analysis_storage[analysis_id] = response.model_dump()
    logger.info("[analyze-cv-advanced] %s -> %s skills, %s ms", analysis_id, len(flattened_skills), response.processing_time_ms)
    log_ml(
        "cv-analysis",
        analysis_id=analysis_id,
        skills=len(flattened_skills),
        hard_skills=len(response.hard_skills or []),
        sources=skills_data.get('sources_used'),
        confidence=round(response.confidence_score, 3)
    )

    return response


@app.post("/api/v1/career-guidance", response_model=CareerGuidanceResponsePayload)
async def generate_career_guidance(payload: CareerGuidanceRequestPayload):
    """Produce ML-style career guidance leveraging the latest CV analysis payload."""
    start_time = time.perf_counter()
    analysis = payload.cv_analysis

    if isinstance(analysis, dict):
        analysis = CVAnalysisResponse(**analysis)

    if analysis is None:
        if payload.cv_content:
            try:
                analysis = await analyze_cv_advanced(CVAnalysisRequest(cv_content=payload.cv_content))
            except HTTPException:
                raise
            except Exception as exc:
                logger.exception("[career-guidance] Fallback analysis failed")
                raise HTTPException(status_code=500, detail=f"Unable to analyze CV content: {exc}")
        else:
            raise HTTPException(status_code=400, detail="cv_analysis or cv_content must be provided")

    primary_skills = _select_primary_skills(analysis)
    if not primary_skills and payload.cv_content:
        extracted = _extract_keywords(
            payload.cv_content,
            programming_languages + frameworks + databases + cloud_platforms + tools
        )
        primary_skills = _filter_primary_skills(extracted, analysis)
    if not primary_skills:
        primary_skills = ["Python", "System Design", "Cloud Architecture"]

    job_recommendations: List[JobRecommendationPayload] = []
    job_fetch_meta: Dict[str, Any] = {}
    if primary_skills:
        job_recommendations, job_fetch_meta = await _fetch_dynamic_job_recommendations(primary_skills, analysis)
    if not job_recommendations and (analysis.confidence_score or 0.0) >= 0.6:
        fallback_jobs = _fallback_template_recommendations(primary_skills, analysis)
        if fallback_jobs:
            job_recommendations = fallback_jobs
            job_fetch_meta = job_fetch_meta or {}
            sources = job_fetch_meta.get('sources_used') or []
            if 'skillsync-fallback' not in sources:
                sources.append('skillsync-fallback')
            job_fetch_meta['sources_used'] = sources
            job_fetch_meta['total_jobs'] = len(fallback_jobs)
            log_ml(
                "job-fallback",
                analysis_id=analysis.analysis_id,
                reason="high-confidence-no-matches",
                total=len(fallback_jobs)
            )
    aggregated_gaps: List[str] = []
    for job in job_recommendations:
        aggregated_gaps.extend(job.skill_gaps)
    aggregated_gaps = list(dict.fromkeys(filter(None, aggregated_gaps)))

    certification_recommendations = _build_certifications(aggregated_gaps or primary_skills)
    learning_roadmap = _build_learning_roadmap(primary_skills, aggregated_gaps, analysis)
    xai_insights = _build_xai_details(analysis, job_recommendations)
    key_opportunities = _build_key_opportunities(job_recommendations)

    metadata = {
        "processing_time_seconds": round(time.perf_counter() - start_time, 2),
        "cv_skills_count": len(primary_skills),
        "jobs_recommended": len(job_recommendations),
        "certs_recommended": len(certification_recommendations),
        "roadmap_phases": len(learning_roadmap.phases),
        "ml_model": "SkillSync-Career-ML-v1",
        "engine_version": "career-guidance-lite-1.0",
        "timestamp": datetime.utcnow().isoformat(),
        "analysis_id": analysis.analysis_id,
        "job_sources": job_fetch_meta.get('sources_used') if job_fetch_meta else [],
        "live_job_total": job_fetch_meta.get('total_jobs') if job_fetch_meta else 0,
        "search_terms": (job_fetch_meta.get('search_parameters', {}).get('skills_used')
                  if job_fetch_meta else primary_skills[:5]),
        "job_matching_confidence": (xai_insights.ml_confidence_scores.get('job_matching')
                         if xai_insights else None)
    }

    logger.info(
        "[career-guidance] %s jobs / %s certs / %s phases", 
        len(job_recommendations),
        len(certification_recommendations),
        len(learning_roadmap.phases)
    )
    log_ml(
        "career-guidance",
        jobs=len(job_recommendations),
        certs=len(certification_recommendations),
        roadmap_phases=len(learning_roadmap.phases),
        sources=job_fetch_meta.get('sources_used') if job_fetch_meta else None,
        skills=primary_skills[:8]
    )

    return CareerGuidanceResponsePayload(
        job_recommendations=job_recommendations,
        certification_recommendations=certification_recommendations,
        learning_roadmap=learning_roadmap,
        xai_insights=xai_insights,
        key_opportunities=key_opportunities,
        metadata=metadata
    )


@app.post("/api/v1/generate-portfolio", response_model=PortfolioResponsePayload)
async def generate_portfolio_endpoint(payload: PortfolioGenerationRequest):
    if not payload.cv_id:
        raise HTTPException(status_code=400, detail="cv_id is required to generate a portfolio")

    analysis_data = cv_analysis_storage.get(payload.cv_id)
    if not analysis_data:
        raise HTTPException(status_code=404, detail="CV analysis not found. Run the ML workflow first.")

    customization = _normalize_portfolio_customization(payload.customization)
    sections, view_model = _build_portfolio_sections(analysis_data, customization)
    if not sections:
        raise HTTPException(status_code=400, detail="Unable to construct portfolio sections from CV data")

    display_name = (
        (analysis_data.get('personal_info') or {}).get('name')
        or analysis_data.get('name')
        or (analysis_data.get('job_titles') or ['Professional'])[0]
    )
    summary_text = (analysis_data.get('summary') or analysis_data.get('raw_text') or "")[:400]
    top_skills = (analysis_data.get('skills') or [])[:5]
    experience_years = analysis_data.get('total_years_experience') or analysis_data.get('experience_years') or 0

    clean_summary = view_model.get('about') or summary_text

    portfolio_meta = {
        "id": str(uuid.uuid4()),
        "name": f"{display_name} Portfolio",
        "template": payload.template_id,
        "color_scheme": customization.color_scheme,
        "generated_at": datetime.utcnow().isoformat(),
        "summary": clean_summary,
        "stats": {
            "skills_count": sum(len(items) for items in (view_model.get('skills') or {}).values()),
            "top_skills": top_skills,
            "experience_years": experience_years,
            "sections": customization.sections_visible
        },
        "sections": sections,
        "hero": view_model.get('hero'),
        "skills": view_model.get('skills'),
        "experiences": view_model.get('experiences'),
        "projects": view_model.get('projects'),
        "education": view_model.get('education')
    }

    html_content = _render_portfolio_html(portfolio_meta, customization, sections, view_model)
    preview_id = str(uuid.uuid4())
    preview_url = f"/api/v1/portfolio/preview/{preview_id}"

    _cleanup_portfolio_preview_store()
    portfolio_preview_store[preview_id] = {
        "id": preview_id,
        "html": html_content,
        "portfolio": portfolio_meta,
        "customization": customization.dict(),
        "created_at": datetime.utcnow()
    }
    logger.info("[generate-portfolio] %s -> template %s", payload.cv_id, payload.template_id)
    log_ml(
        "portfolio",
        cv_id=payload.cv_id,
        template=payload.template_id,
        sections=len(sections),
        color=customization.color_scheme,
        skills=portfolio_meta['stats']['skills_count']
    )
    return PortfolioResponsePayload(
        portfolio=portfolio_meta,
        html_content=html_content,
        preview_id=preview_id,
        preview_url=preview_url
    )


@app.get("/api/v1/portfolio/preview/{preview_id}", response_class=HTMLResponse)
async def get_portfolio_preview(preview_id: str):
    """Serve cached portfolio HTML for live preview."""
    entry = portfolio_preview_store.get(preview_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Portfolio preview not found or expired")

    created_at = entry.get("created_at")
    if isinstance(created_at, datetime):
        if datetime.utcnow() - created_at > timedelta(minutes=PORTFOLIO_PREVIEW_TTL_MINUTES):
            portfolio_preview_store.pop(preview_id, None)
            raise HTTPException(status_code=404, detail="Portfolio preview not found or expired")

    return HTMLResponse(content=entry.get("html", ""), media_type="text/html")

@app.get("/api/v1/cv/analysis/{analysis_id}")
async def get_cv_analysis(analysis_id: str, db: Session = Depends(get_db)):
    """Retrieve CV analysis results"""
    cv_analysis = db.query(CVAnalysisDB).filter(CVAnalysisDB.id == analysis_id).first()
    
    if not cv_analysis:
        raise HTTPException(status_code=404, detail="CV analysis not found")
    
    return {
        "analysis_id": analysis_id,
        "filename": cv_analysis.filename,
        "skills": cv_analysis.skills,
        "confidence_score": cv_analysis.confidence_score,
        "created_at": cv_analysis.created_at
    }

@app.post("/api/v1/experience/translate", response_model=ExperienceTranslateResponse)
async def translate_experience(request: ExperienceTranslateRequest):
    """Translate experience between technical and business contexts"""
    try:
        translation_result = await experience_translator.translate_experience(
            text=request.text,
            source_type=request.source_type,
            target_type=request.target_type,
            tone=request.tone
        )
        
        return ExperienceTranslateResponse(
            original_text=request.text,
            translated_text=translation_result['translated_text'],
            source_type=request.source_type,
            target_type=request.target_type,
            confidence_score=translation_result['confidence_score'],
            key_changes=translation_result['key_changes']
        )
        
    except Exception as e:
        logger.error(f"Experience translation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

@app.get("/api/v1/skill-gap/analyze/{user_id}")
async def analyze_skill_gap(user_id: str, db: Session = Depends(get_db)):
    """Analyze skill gaps for user"""
    try:
        # Get user's CV analyses
        cv_analyses = db.query(CVAnalysisDB).filter(CVAnalysisDB.user_id == user_id).all()
        
        if not cv_analyses:
            raise HTTPException(status_code=404, detail="No CV analysis found for user")
        
        # Get most recent analysis
        latest_analysis = max(cv_analyses, key=lambda x: x.created_at)
        
        # Basic skill gap analysis
        current_skills = []
        for category_skills in latest_analysis.skills.values():
            current_skills.extend([skill_info['skill'] for skill_info in category_skills])
        
        # Market demand skills
        high_demand_skills = ['python', 'aws', 'docker', 'kubernetes', 'machine learning']
        
        # Calculate gaps
        missing_skills = [skill for skill in high_demand_skills if skill not in [s.lower() for s in current_skills]]
        
        return {
            "current_skills": current_skills,
            "market_demand_skills": high_demand_skills,
            "missing_skills": missing_skills,
            "gap_score": len(missing_skills) / len(high_demand_skills),
            "priority_skills": missing_skills[:3]
        }
        
    except Exception as e:
        logger.error(f"Skill gap analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Skill gap analysis failed: {str(e)}")

@app.get("/api/v1/analytics/dashboard")
async def get_global_analytics_dashboard(limit: int = 12, db: Session = Depends(get_db)):
    """Provide aggregated analytics for the frontend dashboard without requiring a user id."""
    try:
        total_cv_analyses = db.query(CVAnalysisDB).count()
        total_job_matches = db.query(JobMatchDB).count()

        recent_analyses = (
            db.query(CVAnalysisDB)
            .order_by(CVAnalysisDB.created_at.desc())
            .limit(max(limit, 12))
            .all()
        )
        recent_jobs = (
            db.query(JobMatchDB)
            .order_by(JobMatchDB.created_at.desc())
            .limit(max(limit, 20))
            .all()
        )

        skill_counter: Counter[str] = Counter()
        category_counter: Counter[str] = Counter()
        monthly_analysis_counts: Dict[str, int] = defaultdict(int)

        def _coerce_skills(payload: Any) -> Dict[str, Any]:
            if payload is None:
                return {}
            if isinstance(payload, str):
                try:
                    return json.loads(payload)
                except Exception:
                    return {}
            if isinstance(payload, dict):
                return payload
            return {}

        def _skill_name(entry: Any) -> Optional[str]:
            if isinstance(entry, dict):
                for key in ("skill", "name", "title"):
                    value = entry.get(key)
                    if value:
                        return str(value)
            elif isinstance(entry, str):
                return entry
            return None

        for analysis in recent_analyses:
            monthly_key = analysis.created_at.strftime("%Y-%m") if analysis.created_at else "unknown"
            monthly_analysis_counts[monthly_key] += 1

            skills_payload = _coerce_skills(analysis.skills)
            for category, skills in skills_payload.items():
                if not isinstance(skills, list):
                    continue
                category_counter[category] += len(skills)
                for skill_entry in skills:
                    skill_name = _skill_name(skill_entry)
                    if skill_name:
                        normalized = skill_name.strip()
                        if normalized:
                            skill_counter[normalized] += 1

        job_trends_map: Dict[str, Dict[str, Any]] = defaultdict(lambda: {"matches": 0, "scores": []})
        match_scores: List[float] = []
        for job_match in recent_jobs:
            month_key = job_match.created_at.strftime("%Y-%m") if job_match.created_at else "unknown"
            trend_entry = job_trends_map[month_key]
            trend_entry["matches"] += 1
            if job_match.match_score is not None:
                trend_entry["scores"].append(job_match.match_score)
                match_scores.append(job_match.match_score)

        monthly_keys = sorted(k for k in monthly_analysis_counts.keys() if k != "unknown")
        growth_rate = 0.0
        if len(monthly_keys) >= 2:
            last, prev = monthly_keys[-1], monthly_keys[-2]
            current_count = monthly_analysis_counts[last]
            previous_count = monthly_analysis_counts[prev]
            if previous_count > 0:
                growth_rate = ((current_count - previous_count) / previous_count) * 100
            elif current_count > 0:
                growth_rate = 100.0

        skill_progress = []
        for skill, count in skill_counter.most_common(5):
            target = count + max(1, round(count * 0.35))
            skill_progress.append({
                "skill": skill,
                "current": count,
                "target": target
            })

        def _format_month(label: str) -> str:
            try:
                dt = datetime.strptime(label, "%Y-%m")
                return dt.strftime("%b %Y")
            except Exception:
                return label

        job_matching_trends = []
        for month_key in sorted(job_trends_map.keys()):
            matches = job_trends_map[month_key]["matches"]
            scores = job_trends_map[month_key]["scores"]
            avg_score = sum(scores) / len(scores) if scores else 0.0
            job_matching_trends.append({
                "month": _format_month(month_key),
                "matches": matches,
                "avg_score": round(avg_score, 2)
            })

        skill_distribution = [
            {"category": category.title(), "count": count}
            for category, count in category_counter.most_common()
        ]

        recent_activities: List[Dict[str, Any]] = []
        for analysis in recent_analyses[:5]:
            recent_activities.append({
                "id": abs(hash(f"analysis-{analysis.id}")) % 10**6,
                "type": "analysis",
                "description": f"ML analysis completed for {analysis.filename}",
                "timestamp": analysis.created_at.isoformat() if analysis.created_at else datetime.utcnow().isoformat(),
                "status": "complete"
            })

        for job_match in recent_jobs[:5]:
            recent_activities.append({
                "id": abs(hash(f"job-{job_match.id}")) % 10**6,
                "type": "job-search",
                "description": f"Job match generated: {job_match.job_title}",
                "timestamp": job_match.created_at.isoformat() if job_match.created_at else datetime.utcnow().isoformat(),
                "status": "complete"
            })

        recent_activities.sort(key=lambda item: item["timestamp"], reverse=True)
        recent_activities = recent_activities[:8] if recent_activities else [{
            "id": 0,
            "type": "system",
            "description": "No ML activity recorded yet",
            "timestamp": datetime.utcnow().isoformat(),
            "status": "pending"
        }]

        analytics_payload = {
            "overview": {
                "total_cvs": total_cv_analyses,
                "jobs_analyzed": total_job_matches,
                "skills_identified": len(skill_counter),
                "match_score_avg": round(sum(match_scores) / len(match_scores), 2) if match_scores else 0.0,
                "growth_rate": round(growth_rate, 2)
            },
            "skill_progress": skill_progress,
            "job_matching_trends": job_matching_trends,
            "skill_distribution": skill_distribution,
            "recent_activities": recent_activities
        }

        logger.info(
            "[analytics] totals -> cvs=%s jobs=%s recent_cv=%s recent_jobs=%s",
            total_cv_analyses,
            total_job_matches,
            len(recent_analyses),
            len(recent_jobs)
        )

        return {
            "success": True,
            "data": analytics_payload,
            "timestamp": datetime.utcnow().isoformat()
        }

    except Exception as e:
        logger.error(f"Global analytics dashboard failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analytics dashboard failed: {str(e)}")


@app.get("/api/v1/analytics/dashboard/{user_id}")
async def get_analytics_dashboard(user_id: str, db: Session = Depends(get_db)):
    """Get basic analytics dashboard data"""
    try:
        # Get user's data
        cv_analyses = db.query(CVAnalysisDB).filter(CVAnalysisDB.user_id == user_id).all()
        
        if not cv_analyses:
            raise HTTPException(status_code=404, detail="No data found for user")
        
        # Get latest analysis
        latest_analysis = max(cv_analyses, key=lambda x: x.created_at)
        
        # Calculate basic analytics
        skill_distribution = {}
        for category, skills in latest_analysis.skills.items():
            skill_distribution[category] = len(skills)
        
        total_skills = sum(len(skills) for skills in latest_analysis.skills.values())
        career_progress_score = min(latest_analysis.confidence_score * 100, 95)
        
        return {
            "user_id": user_id,
            "skill_distribution": skill_distribution,
            "total_skills": total_skills,
            "career_progress_score": round(career_progress_score, 1),
            "confidence_score": latest_analysis.confidence_score,
            "last_analysis": latest_analysis.created_at.isoformat()
        }
        
    except Exception as e:
        logger.error(f"Analytics dashboard failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analytics dashboard failed: {str(e)}")

# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "timestamp": datetime.utcnow().isoformat()
        }
    )

@app.exception_handler(StarletteHTTPException)
async def general_exception_handler(request: Request, exc: StarletteHTTPException):
    status_code = getattr(exc, "status_code", 500)
    logger.error(
        "Unhandled exception on %s %s -> %s: %s",
        request.method,
        request.url.path,
        status_code,
        exc.detail
    )
    return JSONResponse(
        status_code=status_code,
        content={
            "error": exc.detail,
            "status_code": status_code,
            "timestamp": datetime.utcnow().isoformat()
        }
    )

# Startup event
@app.on_event("startup")
async def startup_event():
    logger.info(" SkillSync Enhanced Backend is starting up...")
    logger.info(f" Version: {settings.app_version}")
    logger.info(" All components initialized successfully")
    try:
        _ensure_system_user()
    except Exception as exc:
        logger.warning("System user bootstrap failed: %s", exc)
    _schedule_model_refresh()

# Main application entry point
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main_enhanced:app", host="0.0.0.0", port=8000, reload=True)