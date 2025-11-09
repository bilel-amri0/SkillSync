"""
F1: CV Upload Multi-format (PDF/DOCX parser)
Extract and parse content from CV files
"""

import PyPDF2
import docx
import pdfplumber
import re
from typing import Dict, Any, List
from pathlib import Path
import tempfile
import aiofiles
from fastapi import UploadFile
import logging

logger = logging.getLogger(__name__)

class CVProcessor:
    """Advanced CV processor with multi-format support"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
        
    async def process_upload(self, file: UploadFile) -> Dict[str, Any]:
        """Process uploaded CV file and extract structured data"""
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name
        
        try:
            # Extract text based on file format
            if file.filename.lower().endswith('.pdf'):
                extracted_data = await self._process_pdf(tmp_file_path)
            elif file.filename.lower().endswith(('.docx', '.doc')):
                extracted_data = await self._process_docx(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file format: {file.filename}")
            
            # Add metadata
            extracted_data.update({
                'filename': file.filename,
                'file_size': len(content),
                'processing_status': 'success'
            })
            
            return extracted_data
            
        except Exception as e:
            logger.error(f"Error processing CV {file.filename}: {str(e)}")
            raise
        finally:
            # Clean up temporary file
            Path(tmp_file_path).unlink(missing_ok=True)
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from PDF files"""
        
        extracted_data = {
            'raw_text': '',
            'personal_info': {},
            'sections': {},
            'contact_info': {},
            'metadata': {}
        }
        
        try:
            # Primary extraction with pdfplumber (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                full_text = ''
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + '\n'
                
                extracted_data['raw_text'] = full_text
                extracted_data['metadata']['total_pages'] = len(pdf.pages)
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    full_text = ''
                    for page in pdf_reader.pages:
                        full_text += page.extract_text() + '\n'
                    
                    extracted_data['raw_text'] = full_text
                    extracted_data['metadata']['total_pages'] = len(pdf_reader.pages)
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {str(e2)}")
                raise ValueError(f"Could not extract text from PDF: {str(e2)}")
        
        # Structure the extracted text
        await self._structure_cv_content(extracted_data)
        
        return extracted_data
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from DOCX files"""
        
        extracted_data = {
            'raw_text': '',
            'personal_info': {},
            'sections': {},
            'contact_info': {},
            'metadata': {}
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            extracted_data['raw_text'] = '\n'.join(paragraphs)
            extracted_data['metadata']['total_paragraphs'] = len(paragraphs)
            
            # Extract tables if present
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            if tables_content:
                extracted_data['tables'] = tables_content
        
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
        
        # Structure the extracted text
        await self._structure_cv_content(extracted_data)
        
        return extracted_data
    
    async def _structure_cv_content(self, extracted_data: Dict[str, Any]):
        """Structure raw text into CV sections"""
        
        raw_text = extracted_data['raw_text']
        
        # Extract personal information
        extracted_data['personal_info'] = self._extract_personal_info(raw_text)
        
        # Extract contact information
        extracted_data['contact_info'] = self._extract_contact_info(raw_text)
        
        # Identify and extract sections
        extracted_data['sections'] = self._extract_sections(raw_text)
        
        return extracted_data
    
    def _extract_personal_info(self, text: str) -> Dict[str, str]:
        """Extract personal information like name, title"""
        
        personal_info = {}
        
        lines = text.split('\n')
        
        # Name is usually in the first few lines and is often capitalized or in larger text
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if len(line) > 2 and len(line.split()) <= 4:
                # Likely a name if it's short and at the top
                if not any(char.isdigit() for char in line) and '@' not in line:
                    personal_info['name'] = line
                    break
        
        # Extract professional title
        title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'consultant', 
                         'specialist', 'coordinator', 'director', 'lead']
        
        for line in lines[:10]:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in title_keywords):
                personal_info['title'] = line.strip()
                break
        
        return personal_info
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        
        contact_info = {}
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone extraction
        phone_patterns = [
            r'\+?\d[\d\s\-\(\)]{8,}\d',
            r'\(\d{3}\)\s?\d{3}[\-\s]?\d{4}',
            r'\d{3}[\-\s]?\d{3}[\-\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break
        
        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact_info['linkedin'] = f"https://{linkedin_matches[0]}"
        
        # GitHub extraction
        github_pattern = r'github\.com/[\w\-]+'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact_info['github'] = f"https://{github_matches[0]}"
        
        return contact_info
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different CV sections"""
        
        sections = {}
        
        # Common section headers
        section_patterns = {
            'experience': r'(experience|work\s+history|employment|professional\s+experience)',
            'education': r'(education|academic|qualifications|degrees)',
            'skills': r'(skills|competencies|technical\s+skills|expertise)',
            'projects': r'(projects|portfolio|work\s+samples)',
            'certifications': r'(certifications?|certificates?|licenses?)',
            'languages': r'(languages?|linguistic)',
            'summary': r'(summary|profile|objective|about)',
            'achievements': r'(achievements?|accomplishments?|awards?)',
            'interests': r'(interests?|hobbies|activities)'
        }
        
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_section_header = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    # Save previous section
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content)
                    
                    current_section = section_name
                    section_content = []
                    is_section_header = True
                    break
            
            if not is_section_header and current_section:
                section_content.append(line)
        
        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections