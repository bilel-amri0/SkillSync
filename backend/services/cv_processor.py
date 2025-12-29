"""
CV Processor Service (F1 & F2)
Handles PDF parsing, text extraction, OCR, and skill extraction using NLP.
"""

import re
import logging
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime

# Document processing
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    Image = None  # Fallback for type hints
    logging.warning("OCR not available. Install pdf2image, pytesseract, Pillow")

ADVANCED_PDF_AVAILABLE = PYPDF2_AVAILABLE and OCR_AVAILABLE

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None
    logging.warning("DOCX processing not available. Install python-docx")

# NLP
try:
    import spacy
    from spacy.matcher import PhraseMatcher
    NLP_AVAILABLE = True
except ImportError:
    NLP_AVAILABLE = False
    logging.warning("spaCy not available. Install spacy and download en_core_web_lg")

logger = logging.getLogger(__name__)


@dataclass
class CVSection:
    """Represents a parsed section of the CV"""
    section_type: str
    content: str
    confidence: float = 0.0


@dataclass
class ExtractedSkill:
    """Represents an extracted skill with metadata"""
    skill: str
    category: str
    confidence: float
    context: str = ""
    source: str = "ner"  # ner, taxonomy, pattern


@dataclass
class ParsedCV:
    """Complete parsed CV data structure"""
    raw_text: str
    personal_info: Dict[str, Any]
    sections: Dict[str, str]
    skills: List[ExtractedSkill]
    experience: List[Dict[str, Any]]
    education: List[Dict[str, Any]]
    metadata: Dict[str, Any]


class CVProcessor:
    """
    Production-ready CV Processor handling multi-format parsing,
    OCR fallback, and NLP-based skill extraction.
    """
    
    def __init__(self, spacy_model: str = "en_core_web_lg"):
        """
        Initialize CV Processor with NLP models and skill taxonomies.
        
        Args:
            spacy_model: spaCy model name to load
        """
        self.spacy_model_name = spacy_model
        self.nlp = None
        self.phrase_matcher = None
        
        # Initialize NLP models
        if NLP_AVAILABLE:
            try:
                self.nlp = spacy.load(spacy_model)
                logger.info(f"✅ Loaded spaCy model: {spacy_model}")
            except OSError:
                logger.warning(f"❌ spaCy model {spacy_model} not found. Run: python -m spacy download {spacy_model}")
                self.nlp = None
        
        # Skill taxonomies (expandable)
        self.skill_taxonomy = self._load_skill_taxonomy()
        
        # Section headers patterns
        self.section_patterns = {
            'experience': [
                r'(?i)^(work\s+)?experience',
                r'(?i)^professional\s+experience',
                r'(?i)^employment\s+history',
                r'(?i)^career\s+history'
            ],
            'education': [
                r'(?i)^education',
                r'(?i)^academic\s+background',
                r'(?i)^qualifications'
            ],
            'skills': [
                r'(?i)^(technical\s+)?skills',
                r'(?i)^competencies',
                r'(?i)^expertise'
            ],
            'certifications': [
                r'(?i)^certifications?',
                r'(?i)^licenses?',
                r'(?i)^professional\s+development'
            ],
            'projects': [
                r'(?i)^projects?',
                r'(?i)^portfolio',
                r'(?i)^notable\s+work'
            ]
        }
        
        logger.info("✅ CVProcessor initialized")
    
    def _load_skill_taxonomy(self) -> Dict[str, List[str]]:
        """
        Load skill taxonomies for matching.
        In production, this would load from ESCO/O*NET databases.
        """
        return {
            'programming_languages': [
                'Python', 'JavaScript', 'Java', 'C++', 'C#', 'TypeScript', 'Go', 'Rust',
                'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'SQL'
            ],
            'web_frameworks': [
                'React', 'Angular', 'Vue', 'Django', 'Flask', 'FastAPI', 'Express',
                'Node.js', 'Spring Boot', 'Laravel', 'Ruby on Rails', 'ASP.NET'
            ],
            'databases': [
                'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Elasticsearch',
                'Oracle', 'SQL Server', 'SQLite', 'Cassandra', 'DynamoDB'
            ],
            'cloud_platforms': [
                'AWS', 'Azure', 'Google Cloud', 'GCP', 'Heroku', 'DigitalOcean',
                'CloudFlare', 'Firebase', 'Vercel', 'Netlify'
            ],
            'devops_tools': [
                'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI', 'GitHub Actions',
                'Terraform', 'Ansible', 'CircleCI', 'Travis CI'
            ],
            'data_science': [
                'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch',
                'scikit-learn', 'Pandas', 'NumPy', 'Keras', 'NLP', 'Computer Vision'
            ],
            'soft_skills': [
                'Leadership', 'Communication', 'Problem Solving', 'Teamwork',
                'Project Management', 'Critical Thinking', 'Adaptability', 'Creativity'
            ]
        }
    
    async def process_cv(
        self, 
        file_content: bytes, 
        filename: str,
        extract_skills: bool = True
    ) -> ParsedCV:
        """
        Main entry point: Process CV file and extract all information.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            extract_skills: Whether to run NLP skill extraction
            
        Returns:
            ParsedCV object with all extracted data
            
        Raises:
            ValueError: If file format is unsupported
            RuntimeError: If processing fails
        """
        logger.info(f"📄 Processing CV: {filename}")
        
        try:
            # Step 1: Extract text based on file type
            file_extension = Path(filename).suffix.lower()
            raw_text = await self._extract_text(file_content, file_extension)
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("CV appears to be empty or too short")
            
            # Step 2: Clean and normalize text
            cleaned_text = self._clean_text(raw_text)
            
            # Step 3: Parse sections
            sections = self._parse_sections(cleaned_text)
            
            # Step 4: Extract personal information
            personal_info = self._extract_personal_info(cleaned_text)
            
            # Step 5: Extract structured data
            experience = self._extract_experience(sections.get('experience', ''))
            education = self._extract_education(sections.get('education', ''))
            
            # Step 6: Extract skills (NLP-based)
            skills = []
            if extract_skills and self.nlp:
                skills = await self._extract_skills_nlp(cleaned_text, sections)
            
            # Step 7: Create metadata
            metadata = {
                'filename': filename,
                'file_size': len(file_content),
                'file_type': file_extension,
                'processed_at': datetime.utcnow().isoformat(),
                'text_length': len(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'sections_found': list(sections.keys()),
                'skills_count': len(skills)
            }
            
            parsed_cv = ParsedCV(
                raw_text=raw_text,
                personal_info=personal_info,
                sections=sections,
                skills=skills,
                experience=experience,
                education=education,
                metadata=metadata
            )
            
            logger.info(f"✅ CV processed successfully: {len(skills)} skills extracted")
            return parsed_cv
            
        except Exception as e:
            logger.error(f"❌ Error processing CV {filename}: {str(e)}", exc_info=True)
            raise RuntimeError(f"Failed to process CV: {str(e)}") from e
    
    async def _extract_text(self, file_content: bytes, file_extension: str) -> str:
        """
        Extract text from file based on extension.
        Implements fallback to OCR for scanned PDFs.
        
        Args:
            file_content: Raw file bytes
            file_extension: File extension (.pdf, .docx, .txt)
            
        Returns:
            Extracted text
        """
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_content)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_content)
            elif file_extension == '.txt':
                return file_content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, pdf_content: bytes) -> str:
        """
        Extract text from PDF with OCR fallback for scanned documents.
        
        Args:
            pdf_content: PDF file bytes
            
        Returns:
            Extracted text
        """
        if not PYPDF2_AVAILABLE:
            raise RuntimeError("PyPDF2 library not installed. Install with: pip install PyPDF2")
        
        try:
            # Attempt 1: Direct text extraction with PyPDF2
            import io
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            extracted_text = "\n".join(text_parts)
            
            # Check if we got meaningful text
            if len(extracted_text.strip()) > 100:
                logger.info(f"✅ Extracted text directly from PDF ({len(extracted_text)} chars)")
                return extracted_text
            
            # Attempt 2: OCR fallback for scanned PDFs
            if OCR_AVAILABLE:
                logger.info("📸 PDF appears to be scanned, attempting OCR...")
                return await self._extract_with_ocr(pdf_content)
            else:
                logger.warning("⚠️ OCR not available, returning partial text")
                return extracted_text
            
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {str(e)}")
            if OCR_AVAILABLE:
                try:
                    logger.info("🔄 Attempting OCR fallback...")
                    return await self._extract_with_ocr(pdf_content)
                except Exception as ocr_error:
                    logger.error(f"❌ OCR fallback also failed: {str(ocr_error)}")
                    raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")
            else:
                raise RuntimeError(f"Failed to extract text from PDF and OCR not available: {str(e)}")
    
    async def _extract_with_ocr(self, pdf_content: bytes) -> str:
        """
        Extract text from scanned PDF using OCR (Tesseract).
        
        Args:
            pdf_content: PDF file bytes
            
        Returns:
            OCR-extracted text
        """
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR libraries not available. Install with: pip install pdf2image pytesseract pillow")
        
        try:
            # Convert PDF to images (300 DPI for good quality)
            images = convert_from_bytes(pdf_content, dpi=300)
            
            text_parts = []
            for i, image in enumerate(images):
                logger.info(f"🔍 OCR processing page {i+1}/{len(images)}")
                
                # Preprocess image for better OCR
                processed_image = self._preprocess_image_for_ocr(image)
                
                # Extract text with Tesseract
                page_text = pytesseract.image_to_string(
                    processed_image,
                    lang='eng+fra',  # English + French
                    config='--psm 6'  # Assume uniform block of text
                )
                text_parts.append(page_text)
            
            ocr_text = "\n".join(text_parts)
            logger.info(f"✅ OCR completed: {len(ocr_text)} chars extracted")
            return ocr_text
            
        except Exception as e:
            logger.error(f"❌ OCR failed: {str(e)}")
            raise RuntimeError(f"OCR extraction failed: {str(e)}") from e
    
    def _preprocess_image_for_ocr(self, image: Any) -> Any:
        """
        Preprocess image to improve OCR accuracy.
        
        Args:
            image: PIL Image object
            
        Returns:
            Preprocessed image
        """
        try:
            import cv2
            import numpy as np
            
            # Convert PIL to OpenCV format
            img_array = np.array(image)
            
            # Convert to grayscale
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # Apply adaptive thresholding for better contrast
            thresh = cv2.adaptiveThreshold(
                blurred, 255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                11, 2
            )
            
            # Convert back to PIL Image
            return Image.fromarray(thresh)
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {str(e)}, using original")
            return image
    
    def _extract_from_docx(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            docx_content: DOCX file bytes
            
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        try:
            import io
            doc = docx.Document(io.BytesIO(docx_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"✅ Extracted text from DOCX ({len(extracted_text)} chars)")
            return extracted_text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\-\(\)\[\]\/\:\;\@\#\+\%\&]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _parse_sections(self, text: str) -> Dict[str, str]:
        """
        Parse CV into structured sections.
        
        Args:
            text: Cleaned CV text
            
        Returns:
            Dictionary mapping section names to content
        """
        sections = {}
        lines = text.split('\n')
        
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            section_found = None
            for section_name, patterns in self.section_patterns.items():
                for pattern in patterns:
                    if re.match(pattern, line):
                        section_found = section_name
                        break
                if section_found:
                    break
            
            if section_found:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        logger.info(f"📑 Parsed {len(sections)} sections: {list(sections.keys())}")
        return sections
    
    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """
        Extract personal information (name, email, phone, location).
        
        Args:
            text: CV text
            
        Returns:
            Dictionary with personal info
        """
        info = {}
        
        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        # Extract phone
        phone_pattern = r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]'
        phones = re.findall(phone_pattern, text)
        if phones:
            info['phone'] = phones[0]
        
        # Extract name (heuristic: first line or first capitalized words)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line.split()) <= 4 and line[0].isupper():
                info['name'] = line
                break
        
        return info
    
    def _extract_experience(self, experience_text: str) -> List[Dict[str, Any]]:
        """
        Extract structured experience entries.
        
        Args:
            experience_text: Experience section text
            
        Returns:
            List of experience dictionaries
        """
        experiences = []
        
        # Simple pattern: look for date ranges
        date_pattern = r'(\d{4})\s*[-–—]\s*(\d{4}|Present|Current)'
        matches = re.finditer(date_pattern, experience_text, re.IGNORECASE)
        
        for match in matches:
            start_year = match.group(1)
            end_year = match.group(2)
            
            # Extract context around date (job title, company)
            start_pos = max(0, match.start() - 200)
            end_pos = min(len(experience_text), match.end() + 200)
            context = experience_text[start_pos:end_pos]
            
            experiences.append({
                'start_year': start_year,
                'end_year': end_year,
                'context': context.strip()
            })
        
        return experiences
    
    def _extract_education(self, education_text: str) -> List[Dict[str, Any]]:
        """
        Extract structured education entries.
        
        Args:
            education_text: Education section text
            
        Returns:
            List of education dictionaries
        """
        education = []
        
        # Look for degree keywords
        degree_keywords = [
            'Bachelor', 'Master', 'PhD', 'Doctorate', 'MBA', 'B.S.', 'M.S.',
            'B.A.', 'M.A.', 'Diploma', 'Certificate', 'Degree'
        ]
        
        lines = education_text.split('\n')
        for line in lines:
            for keyword in degree_keywords:
                if keyword.lower() in line.lower():
                    education.append({
                        'degree': line.strip(),
                        'full_text': line
                    })
                    break
        
        return education
    
    async def _extract_skills_nlp(
        self, 
        text: str, 
        sections: Dict[str, str]
    ) -> List[ExtractedSkill]:
        """
        Extract skills using spaCy NER and taxonomy matching.
        
        Args:
            text: Full CV text
            sections: Parsed CV sections
            
        Returns:
            List of ExtractedSkill objects
        """
        if not self.nlp:
            logger.warning("spaCy not available, falling back to pattern matching")
            return self._extract_skills_fallback(text)
        
        skills = []
        seen_skills = set()
        
        # Process with spaCy
        doc = self.nlp(text[:100000])  # Limit to 100K chars for performance
        
        # Method 1: Extract from NER (SKILL, PRODUCT, ORG entities)
        for ent in doc.ents:
            if ent.label_ in ['PRODUCT', 'ORG', 'GPE']:
                skill_text = ent.text.strip()
                if self._is_valid_skill(skill_text) and skill_text not in seen_skills:
                    skills.append(ExtractedSkill(
                        skill=skill_text,
                        category=self._categorize_skill(skill_text),
                        confidence=0.7,
                        context=ent.sent.text[:100],
                        source='ner'
                    ))
                    seen_skills.add(skill_text)
        
        # Method 2: Taxonomy matching
        for category, skill_list in self.skill_taxonomy.items():
            for skill in skill_list:
                # Case-insensitive search
                pattern = re.compile(r'\b' + re.escape(skill) + r'\b', re.IGNORECASE)
                matches = pattern.finditer(text)
                
                for match in matches:
                    if skill not in seen_skills:
                        # Find context sentence
                        start = max(0, match.start() - 50)
                        end = min(len(text), match.end() + 50)
                        context = text[start:end]
                        
                        skills.append(ExtractedSkill(
                            skill=skill,
                            category=category,
                            confidence=0.9,  # High confidence for taxonomy matches
                            context=context,
                            source='taxonomy'
                        ))
                        seen_skills.add(skill)
                        break  # Only count once
        
        # Sort by confidence
        skills.sort(key=lambda x: x.confidence, reverse=True)
        
        logger.info(f"🎯 Extracted {len(skills)} skills using NLP")
        return skills
    
    def _extract_skills_fallback(self, text: str) -> List[ExtractedSkill]:
        """
        Fallback skill extraction without spaCy (pattern matching only).
        
        Args:
            text: CV text
            
        Returns:
            List of ExtractedSkill objects
        """
        skills = []
        seen_skills = set()
        
        for category, skill_list in self.skill_taxonomy.items():
            for skill in skill_list:
                if skill.lower() in text.lower() and skill not in seen_skills:
                    skills.append(ExtractedSkill(
                        skill=skill,
                        category=category,
                        confidence=0.8,
                        context="",
                        source='pattern'
                    ))
                    seen_skills.add(skill)
        
        return skills
    
    def _is_valid_skill(self, text: str) -> bool:
        """
        Validate if extracted text is likely a skill.
        
        Args:
            text: Potential skill text
            
        Returns:
            True if valid skill
        """
        # Filter out common false positives
        invalid_patterns = [
            r'^\d+$',  # Just numbers
            r'^[A-Z]{1}$',  # Single letter
            r'^(the|and|or|but|in|on|at|to|for)$'  # Common words
        ]
        
        for pattern in invalid_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return False
        
        # Valid if 2-50 chars and contains alphanumeric
        return 2 <= len(text) <= 50 and any(c.isalnum() for c in text)
    
    def _categorize_skill(self, skill: str) -> str:
        """
        Categorize a skill based on taxonomy.
        
        Args:
            skill: Skill text
            
        Returns:
            Category name
        """
        for category, skills in self.skill_taxonomy.items():
            if skill in skills:
                return category
        
        return 'other'
